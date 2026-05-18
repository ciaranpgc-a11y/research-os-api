"""Tests for extract extraction service functions (mocked OpenAI)."""

import json
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

import pytest


def _setup_extract_db(monkeypatch, tmp_path):
    """Point the DB at a fresh SQLite file and reset singletons."""
    db_path = tmp_path / "extract_extraction_test.db"
    monkeypatch.setenv("DATABASE_URL", f"sqlite+pysqlite:///{db_path}")
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")
    monkeypatch.setenv("EXTRACT_OPENAI_API_KEY", "test-key")
    monkeypatch.setenv("DATA_LIBRARY_ROOT", str((tmp_path / "dl").resolve()))

    from research_os.db import reset_database_state

    reset_database_state()

    import research_os.extract_auth.service as auth_svc

    auth_svc._admin_seeded = False


@pytest.fixture()
def extraction_svc(monkeypatch, tmp_path):
    _setup_extract_db(monkeypatch, tmp_path)
    from research_os.extract_extraction import service as svc

    return svc


# ---------------------------------------------------------------------------
# Prompt selection tests
# ---------------------------------------------------------------------------


def test_prompt_selection_rhc_email(extraction_svc):
    prompt_file = extraction_svc.PROMPT_MAP.get(("rhc", "email"))
    assert prompt_file == "email_extraction.txt"


def test_prompt_selection_rhc_report(extraction_svc):
    prompt_file = extraction_svc.PROMPT_MAP.get(("rhc", "report"))
    assert prompt_file == "report_extraction.txt"


def test_prompt_selection_rhc_default(extraction_svc):
    # When source_type is None for RHC, _get_prompt defaults to "report"
    prompt_file = extraction_svc.PROMPT_MAP.get(("rhc", None))
    assert prompt_file == "report_extraction.txt"


def test_prompt_selection_echo(extraction_svc):
    prompt_file = extraction_svc.PROMPT_MAP.get(("echo", None))
    assert prompt_file == "echo_extraction.txt"


def test_prompt_selection_cmr(extraction_svc):
    prompt_file = extraction_svc.PROMPT_MAP.get(("cmr", None))
    assert prompt_file == "cmr_extraction.txt"


# ---------------------------------------------------------------------------
# File validation tests
# ---------------------------------------------------------------------------


def test_file_type_validation(extraction_svc):
    """Unsupported file extensions should raise ExtractionUnsupportedFileError."""
    with pytest.raises(extraction_svc.ExtractionUnsupportedFileError):
        extraction_svc.extract_from_file(
            file_bytes=b"MZ\x90\x00",  # fake .exe header
            filename="malware.exe",
            modality="rhc",
        )


def test_file_size_validation(extraction_svc):
    """Files over 20 MB should raise ExtractionFileTooLargeError."""
    oversized = b"\x00" * (20 * 1024 * 1024 + 1)
    with pytest.raises(extraction_svc.ExtractionFileTooLargeError):
        extraction_svc.extract_from_file(
            file_bytes=oversized,
            filename="huge.pdf",
            modality="rhc",
        )


def test_docx_table_text_is_included_for_file_extraction(extraction_svc, monkeypatch):
    """DOCX table cell text should be sent to the model, not just paragraphs."""
    from io import BytesIO

    from docx import Document

    doc = Document()
    doc.add_paragraph("Echo Report")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Study Date"
    table.cell(0, 1).text = "16/04/2026"
    table.cell(1, 0).text = "TAPSE"
    table.cell(1, 1).text = "22 mm"
    buffer = BytesIO()
    doc.save(buffer)

    captured: dict[str, str] = {}

    def fake_extract_from_text(text, modality, source_type=None):
        captured["text"] = text
        captured["modality"] = modality
        return {
            "modality": modality,
            "source_type": source_type,
            "extracted_data": {"tapse": 22},
        }

    monkeypatch.setattr(extraction_svc, "extract_from_text", fake_extract_from_text)

    result = extraction_svc.extract_from_file(buffer.getvalue(), "echo.docx", "echo")

    assert captured["modality"] == "echo"
    assert "Echo Report" in captured["text"]
    assert "Study Date | 16/04/2026" in captured["text"]
    assert "TAPSE | 22 mm" in captured["text"]
    assert result["extracted_data"]["tapse"] == 22


def test_doc_extension_accepts_renamed_docx(extraction_svc, monkeypatch):
    """A DOC extension should still work if the bytes are actually DOCX."""
    from io import BytesIO

    from docx import Document

    doc = Document()
    doc.add_paragraph("Legacy extension, modern Word payload")
    buffer = BytesIO()
    doc.save(buffer)

    captured: dict[str, str] = {}

    def fake_extract_from_text(text, modality, source_type=None):
        captured["text"] = text
        return {
            "modality": modality,
            "source_type": source_type,
            "extracted_data": {"study_date": "16/04/2026"},
        }

    monkeypatch.setattr(extraction_svc, "extract_from_text", fake_extract_from_text)

    result = extraction_svc.extract_from_file(buffer.getvalue(), "echo.doc", "echo")

    assert "Legacy extension, modern Word payload" in captured["text"]
    assert result["extracted_data"]["study_date"] == "16/04/2026"


def test_legacy_doc_uses_converter_for_text_extraction(extraction_svc, monkeypatch):
    """True binary DOC uploads should use the configured converter."""
    captured: dict[str, str] = {}

    def fake_extract_from_text(text, modality, source_type=None):
        captured["text"] = text
        return {
            "modality": modality,
            "source_type": source_type,
            "extracted_data": {"tapse": 22},
        }

    def fake_which(name):
        return "/usr/bin/antiword" if name == "antiword" else None

    def fake_run(args, **kwargs):
        assert args[0] == "/usr/bin/antiword"
        return SimpleNamespace(returncode=0, stdout="Echo report\nTAPSE 22 mm\n", stderr="")

    monkeypatch.setattr(extraction_svc, "extract_from_text", fake_extract_from_text)
    monkeypatch.setattr(extraction_svc.shutil, "which", fake_which)
    monkeypatch.setattr(extraction_svc.subprocess, "run", fake_run)

    result = extraction_svc.extract_from_file(b"\xd0\xcf\x11\xe0legacy-doc", "echo.doc", "echo")

    assert captured["text"] == "Echo report\nTAPSE 22 mm"
    assert result["extracted_data"]["tapse"] == 22


# ---------------------------------------------------------------------------
# Extraction with mocked OpenAI
# ---------------------------------------------------------------------------


def _make_mock_response(extracted_data: dict) -> MagicMock:
    """Build a mock that mimics openai.chat.completions.create() response."""
    mock_message = MagicMock()
    mock_message.content = json.dumps(extracted_data)
    mock_choice = MagicMock()
    mock_choice.message = mock_message
    mock_response = MagicMock()
    mock_response.choices = [mock_choice]
    return mock_response


def test_extraction_with_mock(extraction_svc):
    """Mock openai client, call extract_from_text, verify extracted_data returned."""
    fake_data = {
        "pa_mean": 25.0,
        "pvr_wu": 3.5,
        "pcwp_mean": 12.0,
        "cardiac_output": 5.0,
    }
    mock_response = _make_mock_response(fake_data)

    with patch.object(extraction_svc, "_get_client") as mock_get_client:
        mock_client = MagicMock()
        mock_client.chat.completions.create.return_value = mock_response
        mock_get_client.return_value = mock_client

        result = extraction_svc.extract_from_text(
            text="Sample RHC report text with measurements",
            modality="rhc",
            source_type="report",
        )

    assert result["modality"] == "rhc"
    assert result["source_type"] == "report"
    assert result["extracted_data"]["pa_mean"] == 25.0
    assert result["extracted_data"]["pvr_wu"] == 3.5


def test_extraction_echo_with_mock(extraction_svc):
    """Verify echo extraction works with mocked client."""
    fake_data = {"lvef": 55.0, "tapse": 22.0, "rvsp": 35.0}
    mock_response = _make_mock_response(fake_data)

    with patch.object(extraction_svc, "_get_client") as mock_get_client:
        mock_client = MagicMock()
        mock_client.chat.completions.create.return_value = mock_response
        mock_get_client.return_value = mock_client

        result = extraction_svc.extract_from_text(
            text="Echo report with measurements",
            modality="echo",
        )

    assert result["modality"] == "echo"
    assert result["extracted_data"]["lvef"] == 55.0


def test_echo_aliases_are_normalized(extraction_svc):
    """Echo prompt alias keys should map onto the review and record schema."""
    normalized = extraction_svc._normalize_extracted_payload(
        "echo",
        {
            "heart_rate_bpm": "88 bpm",
            "pulmonary_hypertension_probability": "High",
            "primary_diagnosis": "Pulmonary hypertension",
            "secondary_pathology": "Right ventricular dilatation",
            "conclusion_text_exact": "Pulmonary hypertension.\nRight ventricular dilatation.",
            "conclusion_items": ["Pulmonary hypertension.", "Right ventricular dilatation."],
            "narrative_text": "Technically limited study.",
            "uncertain_fields": ["rvsp"],
            "extraction_warnings": ["RVSP excluded RAP"],
            "ai_confidence": "medium",
        },
    )

    assert normalized["hr"] == 88
    assert normalized["ph_prob"] == "High"
    assert normalized["primary_dx"] == "Pulmonary hypertension"
    assert normalized["secondary_path"] == "Right ventricular dilatation"
    assert normalized["conclusion"] == "Pulmonary hypertension.\nRight ventricular dilatation."
    assert normalized["conc_items"] == "Pulmonary hypertension.\nRight ventricular dilatation."
    assert normalized["narrative"] == "Technically limited study."
    assert normalized["uncertain"] == "rvsp"
    assert normalized["ai_warnings"] == "RVSP excluded RAP"
    assert normalized["ai_conf"] == "medium"


def test_echo_unit_values_are_normalized(extraction_svc):
    """Echo extraction should convert source units into the UI/schema units."""
    normalized = extraction_svc._normalize_extracted_payload(
        "echo",
        {
            "height": "1.78 m",
            "tapse": "2.1 cm",
            "mapse_cm": 1.2,
            "lvidd": "5.1 cm",
            "ivsd": "0.9 cm",
            "main_pulmonary_artery_diameter_cm": "3.2 cm",
            "rv_s_prime_m_s": "0.11 m/s",
            "tr_vmax": "310 cm/s",
            "lvot_vti": "220 mm",
            "dt_ms": "0.18 s",
        },
    )

    assert normalized["height"] == 178
    assert normalized["tapse"] == 21
    assert normalized["mapse"] == 12
    assert normalized["lvidd"] == 51
    assert normalized["ivsd"] == 9
    assert normalized["main_pa_mm"] == 32
    assert normalized["rv_s"] == 11
    assert normalized["tr_vmax"] == 3.1
    assert normalized["lvot_vti"] == 22
    assert normalized["dt_ms"] == 180


def test_echo_text_backfills_measurements_and_classification(extraction_svc):
    """Echo normalization should rescue common table/narrative values missed by GPT."""
    source_text = """
Echo Report
Study Date | 02/04/2025
Reported By | Beatriz Vazquez
Rhythm : Rate | Sinus rhythm 78 bpm
Height | 1.72 m
Weight | 82 kg
LVEF | 62 %
GLS | -18 %
TAPSE | 2.1 cm
RV S' | 11 cm/s
TR Vmax | 3.4 m/s
RVSP | 52 mmHg
RAP | 10 mmHg
LVIDd | 5.1 cm
IVSd | 0.9 cm
LVPWd | 0.9 cm
LA volume index | 38 ml/m2
RV basal diameter | 4.4 cm
Ascending aorta mid | 4.1 cm

Conclusions:
High probability of pulmonary hypertension.
Right ventricular dilatation and dysfunction.
Mild mitral regurgitation.
""".strip()

    normalized = extraction_svc._normalize_extracted_payload(
        "echo",
        {
            "conclusion_items": [
                "High probability of pulmonary hypertension.",
                "Right ventricular dilatation and dysfunction.",
            ],
        },
        source_text,
    )

    assert normalized["study_date"] == "02/04/2025"
    assert normalized["hr"] == 78
    assert normalized["height"] == 172
    assert normalized["weight"] == 82
    assert normalized["bsa"] == 1.98
    assert normalized["lvef"] == 62
    assert normalized["gls"] == -18
    assert normalized["tapse"] == 21
    assert normalized["rv_s"] == 11
    assert normalized["tr_vmax"] == 3.4
    assert normalized["rvsp"] == 52
    assert normalized["rap"] == 10
    assert normalized["lvidd"] == 51
    assert normalized["ivsd"] == 9
    assert normalized["lvpwd"] == 9
    assert normalized["la_voli"] == 38
    assert normalized["rvd1"] == 44
    assert normalized["asc_ao_mid"] == 41
    assert normalized["ph_prob"] == "High"
    assert normalized["case_type"] == "Pulmonary hypertension / right heart"
    assert normalized["primary_dx"] == "Pulmonary hypertension"
    assert normalized["secondary_path"] == "Right ventricular dilatation and dysfunction"
    assert normalized["mr_grade"] == "Mild"


def test_echo_doc_table_backfill_populates_review_form_fields(extraction_svc):
    """Converted Word echo tables should populate visible review fields."""
    source_text = """
|Hospital Number:   |0688491                |Date of     |30/09/1965         |
|                   |                       |Birth:      |                   |
|Consultant:        |Dr M Marquette         |Ward/Op:    |OP                 |
|Date:              |16/11/2024             |Operator:   |Louise Hutchinson  |
|Purpose of study:  |Severe emphysema work up for surgical intervention      |
|Rhythm during      |Sinus     |HR |90     |Image quality |Adequate with    |
|study:             |rhythm    |   |       |              |limited images   |
|                   |          |   |       |              |from some        |
|                   |          |   |       |              |projections      |
|Height (cm)                                                                 |
|Conclusions                                                                 |
|Normal size left ventricle with normal wall thickness and normal systolic   |
|function.                                                                   |
|Normal size right ventricle with normal systolic function.                  |
|Normal size atria.                                                          |
|No significant valvular abnormalities detected.                             |
|Normal female adult dimensions (mm)                                         |
|LVIDd              |43     |35-51    |LVIDs           |31          |20-37    |
|IVSd               |7      |5-11     |LVPWd           |7           |6-12     |
|LVEDV indexed      |27.4   |29-70    |LVESV indexed   |            |8-27     |
|LV Mass (g)        |88.5   |51-173   |LV Mass indexed |52.4        |33-99    |
|LA Vol indexed     |21.4   |<34ml/m2 |Ao @ sinus      |      |      |13.1-20.7|
|RA Area            |7      |<=19cm2  |Sino tubular    |      |      |11.0-17.8|
|RA Area indexed    |4.1    |<=11.0cm2/|Proximal       |30    |17.6  |11.4-19.8|
|                   |       |m2       |ascending Ao    |      |      |         |
|RVD1 base          |28     |22-43    |Right para mid  |      |      |11.4-19.8|
|RVD2 mid           |18     |17-35    |Ao arch         |21    |      |-        |
|RVD3 long          |68     |51-80    |LVOT diameter   |20          |-        |
|RVOT2 distal       |       |14-28    |IVC (mm)        |14          |21       |
""".strip()

    normalized = extraction_svc._normalize_extracted_payload("echo", {}, source_text)

    assert normalized["study_date"] == "16/11/2024"
    assert normalized["consultant"] == "Dr M Marquette"
    assert normalized["ward_op"] == "OP"
    assert normalized["study_reason"] == "Severe emphysema work up for surgical intervention"
    assert normalized["reported_by"] == "Louise Hutchinson"
    assert normalized["rhythm"] == "Sinus rhythm"
    assert normalized["hr"] == 90
    assert normalized["image_quality"] == "Adequate with limited images from some projections"
    assert normalized["lv_size"] == "Normal"
    assert normalized["lv_wall"] == "Normal"
    assert normalized["lv_fn"] == "Normal"
    assert normalized["rv_size"] == "Normal"
    assert normalized["rv_fn"] == "Normal"
    assert normalized["la_size"] == "Normal"
    assert normalized["ra_size"] == "Normal"
    assert normalized["lvidd"] == 43
    assert normalized["lvids"] == 31
    assert normalized["ivsd"] == 7
    assert normalized["lvpwd"] == 7
    assert normalized["lvedvi"] == 27.4
    assert "lvesvi" not in normalized
    assert normalized["lvmi"] == 52.4
    assert normalized["la_voli"] == 21.4
    assert normalized["ra_area"] == 7
    assert normalized["ra_areai"] == 4.1
    assert normalized["rvd1"] == 28
    assert normalized["rvd2"] == 18
    assert normalized["rvd3"] == 68
    assert "rvot2" not in normalized
    assert normalized["asc_ao_prox"] == 30
    assert normalized["ao_arch"] == 21
    assert normalized["lvot_diam"] == 20
    assert normalized["ivc_size"] == "14"
    assert normalized["mr_grade"] == "None"
    assert normalized["case_type"] == "Normal"
    assert normalized["primary_dx"] == "Normal echo"


def test_extraction_cmr_includes_json_instruction(extraction_svc):
    """CMR text extraction must mention JSON when using json_object response_format."""
    fake_data = {"study_date": "2026-04-16", "primary_diagnosis": "Normal"}
    mock_response = _make_mock_response(fake_data)

    with patch.object(extraction_svc, "_get_client") as mock_get_client:
        mock_client = MagicMock()
        mock_client.chat.completions.create.return_value = mock_response
        mock_get_client.return_value = mock_client

        extraction_svc.extract_from_text(
            text="CMR report text",
            modality="cmr",
        )

    kwargs = mock_client.chat.completions.create.call_args.kwargs
    assert kwargs["response_format"] == {"type": "json_object"}
    contents = []
    for message in kwargs["messages"]:
        content = message.get("content", "")
        if isinstance(content, str):
            contents.append(content.lower())
    assert any("json" in content for content in contents)
    assert any("uncertain_fields" in content for content in contents)
    assert any("extraction_warnings" in content for content in contents)


def test_cmr_aliases_are_normalized(extraction_svc):
    """CMR alias keys from the prompt should map onto the extract record schema."""
    normalized = extraction_svc._normalize_extracted_payload(
        "cmr",
        {
            "study_date": "2026-04-16",
            "patient_class": "Pulmonary hypertension / right heart",
            "primary_diagnosis": "Pulmonary hypertension phenotype",
            "secondary_diagnosis": "RV insertion point fibrosis",
            "conclusion_items": ["Pulmonary hypertension phenotype", "RV insertion point fibrosis"],
            "extraction_warnings": ["No explicit stress data"],
            "uncertain_fields": ["flow"],
        },
    )

    assert normalized["date_cmr"] == "16/04/2026"
    assert normalized["cmr_class"] == "Pulmonary hypertension / right heart"
    assert normalized["primary_dx"] == "Pulmonary hypertension phenotype"
    assert normalized["secondary_dx"] == "RV insertion point fibrosis"
    assert normalized["conclusions"] == "Pulmonary hypertension phenotype\nRV insertion point fibrosis"
    assert normalized["qc_notes"] == "No explicit stress data\nUncertain fields: flow"


def test_rhc_height_weight_aliases_and_units_are_normalized(extraction_svc):
    """RHC height/weight should populate numeric review fields even with labels/units."""
    normalized = extraction_svc._normalize_extracted_payload(
        "rhc",
        {
            "Procedure Date": "2026-04-16",
            "Height (cm)": "180 cm",
            "weight_kg": "74 kg",
            "PVR": "3.4 WU",
            "raw_extracted_text": "Height 180 cm, weight 74 kg",
        },
    )

    assert normalized["date_rhc"] == "16/04/2026"
    assert normalized["height"] == 180
    assert normalized["weight"] == 74
    assert normalized["pvr_wu"] == 3.4
    assert normalized["raw_text"] == "Height 180 cm, weight 74 kg"


def test_cmr_text_backfills_and_derivations(extraction_svc):
    """CMR normalization should backfill deterministic values from report text."""
    source_text = """
Patient Name: ROBERT CHAMBERS
Clinical History
severe emphysema for LVR intervention - FAO Dr Garg

Height 180 cm and weight 74 kg (BMI 23). Heart rate was between 62-69 bpm throughout the scan.
The left and right atria are normal in size.
4D flow assessment:
Late gadolinium enhancement: there is a small subendocardial patch of enhancement in the mid septal wall.
Global RV systolic function is preserved (RVEF 57%); however, longitudinal systolic function is reduced (TAPSE 12 mm).
LA volume (mL/m2)                        32
Low probability of significant pulmonary hypertension on CMR.
""".strip()

    normalized = extraction_svc._normalize_extracted_payload(
        "cmr",
        {
            "study_date": "2026-04-16",
            "height": "180",
            "weight": "74",
            "primary_diagnosis": "Non-ischaemic cardiomyopathy",
            "contrast": "Yes",
            "lvedvi": "58",
            "lvesvi": "19",
            "lvsvi": "39",
            "lvmi": "48",
            "rvedvi": "72",
            "rvesvi": "31",
            "rvsvi": "41",
        },
        source_text,
    )

    assert normalized["date_cmr"] == "16/04/2026"
    assert normalized["indication"] == "severe emphysema for LVR intervention - FAO Dr Garg"
    assert normalized["heart_rate"] == "65.5"
    assert normalized["flow"] == "4D-flow"
    assert normalized["contrast"] == "Gadolinium"
    assert normalized["tapse"] == "12"
    assert normalized["la_size"] == "Normal"
    assert normalized["ra_size"] == "Normal"
    assert normalized["lvedv"] == "111.6"
    assert normalized["rvedv"] == "138.5"
    assert normalized["rv_lv_ratio"] == "1.24"
    assert normalized["la_volume"] == "61.6"
    assert normalized["ph"] == "Low"
    assert normalized["primary_dx"] == "Prior myocardial infarction / ischaemic scar"
    assert normalized["cmr_class"] == "Ischaemic"


def test_cmr_text_backfills_full_table_and_narrative(extraction_svc):
    """CMR normalization should recover values from standard quantitative tables."""
    source_text = """
Patient Name: ROBERT CHAMBERS
Date of Birth: 19/05/1953
Local Hospital Number: 0330643

Clinical History

severe emphysema for LVR intervention - FAO Dr Garg

Height 180 cm and weight 74 kg (BMI 23). Heart rate was between 62-69 bpm throughout the scan.
The LV is normal in size with mild eccentric hypertrophy. Global and longitudinal LV systolic function are preserved (LVEF 68%, MAPSE 10 mm).
The RV is not dilated with preserved global RV systolic function (RVEF 57%); however, longitudinal systolic function is reduced (TAPSE 12 mm).
The left and right atria are normal in size. The interatrial septum bows towards the left atrium, consistent with elevated right atrial pressure.
There is mild mitral regurgitation. No significant aortic or tricuspid regurgitation.

CMR quantitative                       Value             Normal range
LV EDV (i) (mL/m2)                       58          70 (44-95)
LV ESV (i) (mL/m2)                       19          24 (10-38)
LV mass (i) (g/m2)                       48          64 (44-84)
LV peak wall thickness (mm)              12           9 (6-12)
LV SV (i) (mL/m2)                        39          46 (28-64)
LV EF (%)                                68          66 (53-79)
RV EDV (i) (mL/m2)                       72          83 (49-117)
RV ESV (i) (mL/m2)                       31          34 (12-56)
RV SV (i) (mL/m2)                        41          47 (29-65)
RV EF (%)                                57          59 (48-71)
LA volume (mL/m2)                        32          36 (13-59)
RA volume (mL/m2)                        27          44 (17-71)
Estimated PCWP (mmHg)                    13             (5-12)

4D flow assessment:
Aortic forward flow (mL/beat)            66
Aortic backward flow (mL/beat)           1
Pulmonary forward flow (mL/beat)         77
Vortex formation is present in the main pulmonary artery (20% of cardiac phases).
MPA systolic diameter (mm)               34          27 (22-33)

Late gadolinium enhancement: there is a small subendocardial patch of enhancement in the mid septal wall.
Native myocardial T1 (ms)               1110            (920-1050)
ECV (%)                                  36             (22-33)
Native myocardial T2 (ms)                53             (46-58)

Low probability of significant pulmonary hypertension on CMR: mildly dilated MPA with vortex formation and reduced TAPSE.

Conclusions:
1. Preserved LV systolic function (LVEF 68%) with a non-dilated LV and no regional wall motion abnormality.
2. Normal RV size with preserved RV systolic function (RVEF 57%).
3. There is evidence of small septal infarct
4. No CMR features to support PH physiology more than mild (at most) severity.
""".strip()

    normalized = extraction_svc._normalize_extracted_payload(
        "cmr",
        {
            "study_date": "30 December 2025",
            "conclusion_items": [
                "Preserved LV systolic function (LVEF 68%) with a non-dilated LV and no regional wall motion abnormality.",
                "Normal RV size with preserved RV systolic function (RVEF 57%).",
                "There is evidence of small septal infarct",
                "No CMR features to support PH physiology more than mild (at most) severity.",
            ],
        },
        source_text,
    )

    assert normalized["date_cmr"] == "30/12/2025"
    assert normalized["indication"] == "severe emphysema for LVR intervention - FAO Dr Garg"
    assert normalized["height"] == "180"
    assert normalized["weight"] == "74"
    assert normalized["heart_rate"] == "65.5"
    assert normalized["contrast"] == "Gadolinium"
    assert normalized["flow"] == "4D-flow"
    assert normalized["lv_size"] == "Normal"
    assert normalized["lv_function"] == "Normal"
    assert normalized["lvef"] == "68"
    assert normalized["lvedvi"] == "58"
    assert normalized["lvesvi"] == "19"
    assert normalized["lvsvi"] == "39"
    assert normalized["lvmi"] == "48"
    assert normalized["max_lv_wall"] == "12"
    assert normalized["rvef"] == "57"
    assert normalized["rvedvi"] == "72"
    assert normalized["rvesvi"] == "31"
    assert normalized["rvsvi"] == "41"
    assert normalized["tapse"] == "12"
    assert normalized["rv_lv_ratio"] == "1.24"
    assert normalized["la_size"] == "Normal"
    assert normalized["ra_size"] == "Normal"
    assert normalized["pcwp"] == "13"
    assert normalized["ao_forward_volume"] == "66"
    assert normalized["ao_backward_volume"] == "1"
    assert normalized["pulmonary_forward_volume"] == "77"
    assert normalized["mpa_size"] == "34"
    assert normalized["mpa_vortex"] == "Present (20% of cardiac phases)"
    assert normalized["native_t1"] == "1110"
    assert normalized["ecv"] == "36"
    assert normalized["t2"] == "53"
    assert normalized["lge"] == "Present"
    assert normalized["lge_pattern"] == "Subendocardial"
    assert normalized["mr_severity"] == "Mild"
    assert normalized["ar_severity"] == "None"
    assert normalized["tr_severity"] == "None"
    assert normalized["ph"] == "Low"
    assert normalized["primary_dx"] == "Prior myocardial infarction / ischaemic scar"
    assert normalized["cmr_class"] == "Ischaemic"


def test_cmr_text_overrides_bad_model_values_from_indexed_table(extraction_svc):
    """Formal indexed CMR table values should correct bad model-filled values."""
    source_text = """
Clinical History
severe emphysema for LVR FAO Dr Garg please

The vascular arrangement is normal. Height 180 cm and weight 63 kg (BMI 20). The heart rate throughout the scan averaged 65 bpm.
The RV is not dilated, with borderline reduced global systolic function (RVEF 45%) and markedly reduced longitudinal function (TAPSE 8 mm).
The left and right atria are normal in size.

CMR quantitative                       Value             Normal range
LV EDV (i) (mL/m2)                       56          74 (46-101)
LV ESV (i) (mL/m2)                       19          25 (10-41)
LV mass (i) (g/m2)                       58          55 (39-72)
LV peak wall thickness (mm)              13           9 (6-12)
LV EF (%)                                67          66 (52-79)
LV SV (i) (mL/m2)                        38          48 (30-67)
RV EDV (i) (mL/m2)                       68          79 (45-114)
RV ESV (i) (mL/m2)                       38          34 (16-52)
RV EF (%)                                45          57 (41-72)
RV SV (i) (mL/m2)                        31          44 (20-69)
LA max volume (i) (mL/m2)                33          37 (20-54)
RA max volume (i) (mL/m2)                19          38 (11-65)
MPA systolic diameter (mm)               33          28 (22-34)
Estimated PCWP (mmHg)                    14           8 (5-12)
Estimated RAP (mmHg)                     8            5 (1-8)

Focal non-ischaemic mid-wall LGE in the basal inferior wall. Separate RV insertion point LGE with mild anterior extension into the adjacent mid septum, most in keeping with pressure-loading related insertion site fibrosis.
Native T1 (ms)                           1033        985 (920-1050)
ECV (%)                                  31          27 (22-33)
Native T2 (ms)                           50          52 (46-65)

Flow (2D-PC)                             Aorta       Pulmonary
Forward flow                             85 mL       80 mL
Backward flow                            2 mL        1 mL
Regurgitant fraction                     2%          1%

CMR features support PH physiology, with subtle systolic septal flattening, markedly reduced longitudinal RV shortening (TAPSE 8 mm), borderline reduced global RV systolic function (RVEF 45%), RV insertion point fibrosis, and persistent disorganized vertical and helical flow in the main pulmonary artery.

Conclusions:
1. Preserved LV systolic function (LVEF 67%) and normal size.
2. Borderline reduced RV systolic function (RVEF 45%) with a non-dilated RV.
3. Focal non-ischaemic mid-wall enhancement in the basal inferior wall alongside RV-insertion point fibrosis with mid-wall extension.
4. CMR findings support pulmonary hypertension physiology with early maladaptive RV response/RV-pulmonary arterial uncoupling.
""".strip()

    normalized = extraction_svc._normalize_extracted_payload(
        "cmr",
        {
            "rvedv": "88",
            "rvesv": "48",
            "rvsv": "40",
            "rvedvi": "48",
            "rvesvi": "26",
            "rvsvi": "22",
            "rv_lv_ratio": "1.55",
        },
        source_text,
    )

    assert normalized["heart_rate"] == "65"
    assert normalized["flow"] == "2D-flow"
    assert normalized["rvedvi"] == "68"
    assert normalized["rvesvi"] == "38"
    assert normalized["rvsvi"] == "31"
    assert normalized["rvedv"] == "120.7"
    assert normalized["rvesv"] == "67.4"
    assert normalized["rvsv"] == "55"
    assert normalized["rv_lv_ratio"] == "1.21"
    assert normalized["la_volume"] == "58.6"
    assert normalized["rap"] == "8"
    assert normalized["ao_forward_volume"] == "85"
    assert normalized["pulmonary_forward_volume"] == "80"
    assert normalized["ar_rf"] == "2"
    assert normalized["pr_rf"] == "1"
    assert normalized["lge_pattern"] == "Mid-wall"
    assert normalized["septal_flattening"] == "Present"
    assert normalized["flattening_phase"] == "Systolic"
    assert normalized["ph"] == "High"
    assert normalized["primary_dx"] == "Pulmonary hypertension phenotype"
    assert normalized["secondary_dx"] == "Non-ischaemic myocardial fibrosis / scar"
    assert normalized["cmr_class"] == "Pulmonary hypertension / right heart"


def test_cmr_does_not_overcall_ischaemia_from_rule_out_history(extraction_svc):
    """Clinical-history rule-out wording should not force an ischaemic primary diagnosis."""
    source_text = """
Patient was admitted under cardiology due to chest pain.
Prior MRI to rule out any infarction/inflammation.
Conclusions:
Left ventricle is not-dilated and the LV systolic function is normal (LVEF 68%).
There is no evidence of any scar or fibrosis on the ventricle.
There is no obvious evidence of myocarditis - this could be due to early resolution.
""".strip()

    normalized = extraction_svc._normalize_extracted_payload(
        "cmr",
        {
            "patient_class": "Non-ischaemic cardiomyopathy",
            "primary_diagnosis": "Non-ischaemic cardiomyopathy",
            "conclusions": (
                "Left ventricle is not-dilated and the LV systolic function is normal (LVEF 68%). "
                "There is no evidence of any scar or fibrosis on the ventricle. "
                "There is no obvious evidence of myocarditis - this could be due to early resolution."
            ),
        },
        source_text,
    )

    assert normalized["primary_dx"] == "Non-ischaemic cardiomyopathy"
    assert "ischaemic" not in normalized["qc_notes"].lower()


def test_cmr_qc_notes_fallback_flags_missing_indication(extraction_svc):
    """QC notes should provide a useful fallback when helper warnings are absent."""
    normalized = extraction_svc._normalize_extracted_payload(
        "cmr",
        {
            "study_date": "2026-04-16",
            "patient_class": "Normal",
            "primary_diagnosis": "Normal",
        },
        "No explicit indication heading here.",
    )

    assert normalized["date_cmr"] == "16/04/2026"
    assert normalized["qc_notes"] == "Review extraction: indication was not clearly captured."


# ---------------------------------------------------------------------------
# Save extraction (integration with patients + records, still mocked OpenAI)
# ---------------------------------------------------------------------------


def test_save_extraction(extraction_svc, monkeypatch, tmp_path):
    """Mock extraction + save, verify patient and record created."""
    from research_os.extract_patients import service as patients_svc
    from research_os.extract_records import service as records_svc

    result = extraction_svc.save_extraction(
        modality="rhc",
        hospital_number="HN500",
        create_patient_if_missing=True,
        patient_data={"name": "Save Test Patient"},
        record_data={"pa_mean": 30.0, "pvr_wu": 5.0, "height_cm": "180 cm", "weight_kg": "74 kg"},
    )

    assert "patient" in result
    assert "record" in result
    assert result["patient"]["hn"] == "HN500"
    assert result["patient"]["name"] == "Save Test Patient"
    assert result["record"]["pa_mean"] == 30.0
    assert result["record"]["height"] == 180
    assert result["record"]["weight"] == 74
    assert result["record"]["status"] == "Completed"
    assert result["record"]["status_date"]

    # Verify patient actually persisted
    patient = patients_svc.get_patient("HN500")
    assert patient["name"] == "Save Test Patient"
    assert patient["rhc_count"] == 1


def test_save_extraction_existing_patient(extraction_svc, monkeypatch, tmp_path):
    """Save extraction for an existing patient without create_patient_if_missing."""
    from research_os.extract_patients import service as patients_svc

    # Pre-create the patient
    patients_svc.create_patient(hn="HN501", name="Existing Patient")

    result = extraction_svc.save_extraction(
        modality="echo",
        hospital_number="HN501",
        create_patient_if_missing=False,
        patient_data={},
        record_data={
            "lvef": 60.0,
            "tapse": 24.0,
            "primary_diagnosis": "Pulmonary hypertension",
            "secondary_pathology": "Right ventricular dilatation",
            "conclusion_items": ["Pulmonary hypertension.", "Right ventricular dilatation."],
        },
    )

    assert result["patient"]["hn"] == "HN501"
    assert result["record"]["lvef"] == 60.0
    assert result["record"]["primary_dx"] == "Pulmonary hypertension"
    assert result["record"]["secondary_path"] == "Right ventricular dilatation"
    assert result["record"]["conc_items"] == "Pulmonary hypertension.\nRight ventricular dilatation."
    assert result["record"]["status"] == "Completed"
    assert result["record"]["status_date"]


def test_save_extraction_coerces_echo_numeric_columns(extraction_svc, monkeypatch, tmp_path):
    """Echo save should tolerate descriptive labels in numeric DB columns."""
    from research_os.extract_patients import service as patients_svc

    patients_svc.create_patient(hn="HN503", name="Echo Numeric Patient")

    result = extraction_svc.save_extraction(
        modality="echo",
        hospital_number="HN503",
        create_patient_if_missing=False,
        patient_data={},
        record_data={
            "lvef": "Normal",
            "rvsp": "42 mmHg",
            "septal_flattening_present": "No",
            "primary_diagnosis": "Pulmonary hypertension",
        },
    )

    assert result["record"]["lvef"] is None
    assert result["record"]["rvsp"] == 42.0
    assert result["record"]["sept_flat"] == 0
    assert result["record"]["primary_dx"] == "Pulmonary hypertension"
    assert result["record"]["status"] == "Completed"


def test_save_extraction_links_uploaded_source_file(extraction_svc, monkeypatch, tmp_path):
    """Uploaded source files should be linked to the saved investigation record."""
    from research_os.extract_patients import service as patients_svc
    from research_os.extract_source_files import service as source_files_svc

    patients_svc.create_patient(hn="HN504", name="Echo File Patient")
    upload = source_files_svc.create_source_file(
        modality="echo",
        filename="echo-report.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=b"example echo report bytes",
    )

    result = extraction_svc.save_extraction(
        modality="echo",
        hospital_number="HN504",
        create_patient_if_missing=False,
        patient_data={},
        record_data={"lvef": 60.0},
        source_file_upload_id=upload["id"],
    )

    linked = source_files_svc.get_source_file(upload["id"])
    assert result["record"]["source_file"] == "echo-report.docx"
    assert result["record"]["source_file_upload"]["id"] == upload["id"]
    assert result["record"]["source_file_upload"]["record_id"] == result["record"]["id"]
    assert linked["hn"] == "HN504"
    assert linked["record_id"] == result["record"]["id"]
    assert linked["byte_size"] == len(b"example echo report bytes")


def test_source_file_word_pdf_preview_requires_converter(extraction_svc, monkeypatch):
    """Word source preview should give a clear error when LibreOffice is unavailable."""
    from research_os.extract_source_files import service as source_files_svc

    upload = source_files_svc.create_source_file(
        modality="echo",
        filename="echo-report.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=b"example echo report bytes",
    )
    monkeypatch.setattr(source_files_svc, "_find_office_converter", lambda: None)

    with pytest.raises(source_files_svc.ExtractSourceFileConversionUnavailableError):
        source_files_svc.get_source_file_pdf_preview(upload["id"])


def test_source_file_word_pdf_preview_converts_to_pdf(extraction_svc, monkeypatch):
    """Word source preview should return generated PDF bytes and PDF metadata."""
    from pathlib import Path

    from research_os.extract_source_files import service as source_files_svc

    upload = source_files_svc.create_source_file(
        modality="cmr",
        filename="cmr-report.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=b"example cmr report bytes",
    )

    monkeypatch.setattr(source_files_svc, "_find_office_converter", lambda: "soffice")

    def fake_run(args, capture_output, text, timeout, check):
        out_dir = Path(args[args.index("--outdir") + 1])
        (out_dir / "source.pdf").write_bytes(b"%PDF-1.4 fake")
        return SimpleNamespace(returncode=0, stderr="", stdout="")

    monkeypatch.setattr(source_files_svc.subprocess, "run", fake_run)

    meta, content = source_files_svc.get_source_file_pdf_preview(upload["id"])

    assert meta["filename"] == "cmr-report.pdf"
    assert meta["content_type"] == "application/pdf"
    assert meta["byte_size"] == len(b"%PDF-1.4 fake")
    assert content == b"%PDF-1.4 fake"


def test_save_extraction_coerces_cmr_complex_values(extraction_svc, monkeypatch, tmp_path):
    """CMR extraction saves should coerce list/dict values into DB-safe text."""
    result = extraction_svc.save_extraction(
        modality="cmr",
        hospital_number="HN502",
        create_patient_if_missing=True,
        patient_data={"name": "Complex Value Patient"},
        record_data={
            "primary_dx": "Prior myocardial infarction / ischaemic scar",
            "conclusions": [
                "Left ventricle is not dilated and LV systolic function is normal.",
                "Right ventricle is not dilated and RV systolic function is preserved.",
            ],
            "perfusion_coronary_territory": {"territory": "LAD"},
        },
    )

    assert result["patient"]["hn"] == "HN502"
    assert result["record"]["status"] == "Completed"
    assert (
        result["record"]["conclusions"]
        == "Left ventricle is not dilated and LV systolic function is normal.\n"
        "Right ventricle is not dilated and RV systolic function is preserved."
    )
    assert result["record"]["perfusion_coronary_territory"] == '{"territory": "LAD"}'
