"""GPT-4o extraction service for structured cardiac measurement extraction."""

from __future__ import annotations

import base64
import json
import logging
import math
import os
import re
import shutil
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import openai
from sqlalchemy import Float, Integer

from research_os.config import get_openai_api_key

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB

VALID_MODALITIES = ("rhc", "echo", "cmr", "generic")

PROMPT_MAP: dict[tuple[str, str | None], str] = {
    ("rhc", "email"): "email_extraction.txt",
    ("rhc", "report"): "report_extraction.txt",
    ("rhc", None): "report_extraction.txt",  # default for RHC
    ("echo", None): "echo_extraction.txt",
    ("cmr", None): "cmr_extraction.txt",
    ("generic", None): "generic_extraction.txt",
}

SUPPORTED_FILE_EXTENSIONS = {".pdf", ".doc", ".docx", ".png", ".jpg", ".jpeg"}
EXTRACTION_SCHEMA_EXCLUDED_FIELDS = {
    "id",
    "hn",
    "created_at",
    "status",
    "status_date",
    "measurements_json",
    "pending",
}
EXTRACTION_RESPONSE_META_FIELDS = {
    "hospital_number",
    "patient_name",
    "date_of_birth",
    "sex",
    "gender",
    "conclusion_items",
    "uncertain_fields",
    "extraction_warnings",
}
RHC_NUMERIC_FIELDS = {
    "height",
    "weight",
    "ra_mean",
    "ra_a",
    "ra_v",
    "ra_o2_sat",
    "rv_systolic",
    "rv_diastolic",
    "rv_mean",
    "rv_o2_sat",
    "pa_systolic",
    "pa_diastolic",
    "pa_mean",
    "pa_o2_sat",
    "pcwp_mean",
    "pcwp_a",
    "pcwp_v",
    "pcwp_o2_sat",
    "aorta_systolic",
    "aorta_diastolic",
    "aorta_mean",
    "aorta_o2_sat",
    "lv_systolic",
    "lv_diastolic",
    "lv_mean",
    "lv_o2_sat",
    "cardiac_output",
    "cardiac_index",
    "pvr_wu",
    "pvr_dyn",
    "tpg",
}
RHC_ALIAS_MAP = {
    "date": "date_rhc",
    "rhc_date": "date_rhc",
    "date_rhc": "date_rhc",
    "procedure_date": "date_rhc",
    "study_date": "date_rhc",
    "date_of_rhc": "date_rhc",
    "height": "height",
    "height_cm": "height",
    "height_m": "height",
    "height_metres": "height",
    "height_meters": "height",
    "patient_height": "height",
    "weight": "weight",
    "weight_kg": "weight",
    "patient_weight": "weight",
    "pvr": "pvr_wu",
    "pvr_wu": "pvr_wu",
    "raw_extracted_text": "raw_text",
    "source_text": "raw_text",
    "rhc_comment": "rhc_comments",
    "rhc_comments": "rhc_comments",
    "comments": "rhc_comments",
}
ECHO_ALIAS_MAP = {
    "date": "study_date",
    "echo_date": "study_date",
    "heart_rate": "hr",
    "heart_rate_bpm": "hr",
    "height_cm": "height",
    "weight_kg": "weight",
    "bsa_m2": "bsa",
    "ward_or_op": "ward_op",
    "lv_size_description": "lv_size",
    "lv_wall_thickness_description": "lv_wall",
    "lv_systolic_function_description": "lv_fn",
    "lvef_percent": "lvef",
    "gls_percent": "gls",
    "rv_size_description": "rv_size",
    "rv_function_description": "rv_fn",
    "rv_s_prime_cm_s": "rv_s",
    "rvsp_mmhg": "rvsp",
    "rap_mmhg": "rap",
    "ivc_size_description": "ivc_size",
    "ivc_collapse_description": "ivc_coll",
    "pulmonary_hypertension_probability": "ph_prob",
    "septal_flattening_present": "sept_flat",
    "septal_bounce_present": "sept_bounce",
    "d_shaped_lv_present": "d_shaped_lv",
    "pulmonary_artery_dilated_present": "pa_dilated",
    "pericardial_effusion_present": "peric_eff",
    "interatrial_septum_intact": "ias_intact",
    "left_atrium_size_description": "la_size",
    "right_atrium_size_description": "ra_size",
    "mitral_valve_description": "mv_desc",
    "tricuspid_valve_description": "tv_desc",
    "aortic_valve_description": "av_desc",
    "pulmonary_valve_description": "pv_desc",
    "aortic_root_description": "ao_root_desc",
    "ascending_aorta_description": "asc_ao_desc",
    "main_pulmonary_artery_diameter_mm": "main_pa_mm",
    "primary_diagnosis": "primary_dx",
    "secondary_diagnosis": "secondary_path",
    "secondary_pathology": "secondary_path",
    "conclusion_text_exact": "conclusion",
    "conclusion_items": "conc_items",
    "narrative_text": "narrative",
    "measurement_table": "meas_table",
    "uncertain_fields": "uncertain",
    "extraction_warnings": "ai_warnings",
    "raw_extracted_text": "ai_raw_text",
    "ai_confidence": "ai_conf",
}
ECHO_MM_FIELDS = {
    "tapse",
    "mapse",
    "lvidd",
    "lvids",
    "ivsd",
    "lvpwd",
    "rvd1",
    "rvd2",
    "rvd3",
    "rvot2",
    "la_diam",
    "ao_ann",
    "ao_sinus",
    "stj_mm",
    "asc_ao_prox",
    "asc_ao_mid",
    "ao_arch",
    "main_pa_mm",
    "lvot_diam",
}
ECHO_SMALL_MOTION_MM_FIELDS = {"tapse", "mapse"}
ECHO_WALL_THICKNESS_MM_FIELDS = {"ivsd", "lvpwd"}
ECHO_CM_PER_SEC_FIELDS = {"rv_s", "med_s", "lat_s", "sept_e", "lat_e", "mv_e", "mv_a"}
ECHO_M_PER_SEC_FIELDS = {"tr_vmax", "av_vmax", "pv_vmax", "lvot_vel"}
ECHO_CM_FIELDS = {"lvot_vti", "av_vti"}
ECHO_MS_FIELDS = {"dt_ms", "ar_pht", "pat"}
ECHO_NUMERIC_FIELDS = (
    ECHO_MM_FIELDS
    | ECHO_CM_PER_SEC_FIELDS
    | ECHO_M_PER_SEC_FIELDS
    | ECHO_CM_FIELDS
    | ECHO_MS_FIELDS
    | {
        "hr",
        "height",
        "weight",
        "bsa",
        "lvef",
        "gls",
        "lvh",
        "rwma",
        "avg_e_ep",
        "sept_e_ep",
        "e_a",
        "fac",
        "rvsp",
        "rap",
        "sept_flat",
        "sept_bounce",
        "d_shaped_lv",
        "pa_dilated",
        "peric_eff",
        "ias_intact",
        "shunt",
        "av_pk_grad",
        "av_mn_grad",
        "ava",
        "la_vol",
        "la_voli",
        "ra_area",
        "ra_areai",
        "lvedvi",
        "lvesvi",
        "rwt",
        "lvmi",
        "dvi",
    }
)
ECHO_UNIT_ALIAS_MAP: dict[str, tuple[str, str]] = {
    "tapse_cm": ("tapse", "cm"),
    "tapse_mm": ("tapse", "mm"),
    "mapse_cm": ("mapse", "cm"),
    "mapse_mm": ("mapse", "mm"),
    "lvidd_cm": ("lvidd", "cm"),
    "lvids_cm": ("lvids", "cm"),
    "ivsd_cm": ("ivsd", "cm"),
    "lvpwd_cm": ("lvpwd", "cm"),
    "rvd1_cm": ("rvd1", "cm"),
    "rvd2_cm": ("rvd2", "cm"),
    "rvd3_cm": ("rvd3", "cm"),
    "rvot2_cm": ("rvot2", "cm"),
    "la_diam_cm": ("la_diam", "cm"),
    "la_diameter_cm": ("la_diam", "cm"),
    "ao_ann_cm": ("ao_ann", "cm"),
    "aortic_annulus_cm": ("ao_ann", "cm"),
    "ao_sinus_cm": ("ao_sinus", "cm"),
    "aortic_sinus_cm": ("ao_sinus", "cm"),
    "stj_cm": ("stj_mm", "cm"),
    "sinotubular_junction_cm": ("stj_mm", "cm"),
    "asc_ao_prox_cm": ("asc_ao_prox", "cm"),
    "ascending_aorta_proximal_cm": ("asc_ao_prox", "cm"),
    "asc_ao_mid_cm": ("asc_ao_mid", "cm"),
    "ascending_aorta_mid_cm": ("asc_ao_mid", "cm"),
    "ao_arch_cm": ("ao_arch", "cm"),
    "aortic_arch_cm": ("ao_arch", "cm"),
    "lvot_diam_cm": ("lvot_diam", "cm"),
    "lvot_diameter_cm": ("lvot_diam", "cm"),
    "main_pa_cm": ("main_pa_mm", "cm"),
    "main_pulmonary_artery_diameter_cm": ("main_pa_mm", "cm"),
    "main_pulmonary_artery_diameter_mm": ("main_pa_mm", "mm"),
    "rv_s_prime_cm_s": ("rv_s", "cm/s"),
    "rv_s_prime_m_s": ("rv_s", "m/s"),
    "rv_s_cm_s": ("rv_s", "cm/s"),
    "rv_s_m_s": ("rv_s", "m/s"),
    "tr_vmax_cm_s": ("tr_vmax", "cm/s"),
    "tr_vmax_m_s": ("tr_vmax", "m/s"),
    "av_vmax_cm_s": ("av_vmax", "cm/s"),
    "av_vmax_m_s": ("av_vmax", "m/s"),
    "pv_vmax_cm_s": ("pv_vmax", "cm/s"),
    "pv_vmax_m_s": ("pv_vmax", "m/s"),
    "lvot_vel_cm_s": ("lvot_vel", "cm/s"),
    "lvot_vel_m_s": ("lvot_vel", "m/s"),
}

# ---------------------------------------------------------------------------
# Prompt loading
# ---------------------------------------------------------------------------


def _get_prompt(modality: str, source_type: str | None) -> str:
    """Load prompt text from the appropriate file."""
    # For echo/cmr/generic, source_type is ignored.
    # For rhc, source_type defaults to "report" if not provided.
    if modality == "rhc":
        key = (modality, source_type or "report")
    else:
        key = (modality, None)

    filename = PROMPT_MAP.get(key)
    if not filename:
        raise ValueError(f"Unknown modality/source_type: {modality}/{source_type}")

    prompt_path = Path(__file__).parent / "prompts" / filename
    return prompt_path.read_text(encoding="utf-8")


# ---------------------------------------------------------------------------
# OpenAI client helpers
# ---------------------------------------------------------------------------


def _get_client() -> openai.OpenAI:
    api_key = os.getenv("EXTRACT_OPENAI_API_KEY") or get_openai_api_key()
    return openai.OpenAI(api_key=api_key, timeout=120.0)


def _get_model() -> str:
    return os.getenv("EXTRACT_OPENAI_MODEL", "gpt-4o")


def _parse_json_response(raw: str) -> dict[str, Any]:
    """Parse a JSON response string, stripping markdown fences if present."""
    text = raw.strip()
    if text.startswith("```"):
        # Strip ```json ... ``` wrappers
        lines = text.split("\n")
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        text = "\n".join(lines)
    return json.loads(text)


def _response_fields_for_modality(modality: str) -> list[str]:
    """Return the extractable field names for a modality."""
    if modality == "generic":
        return []

    from research_os.extract_records import service as records_svc

    fields = records_svc.get_valid_columns(modality) - EXTRACTION_SCHEMA_EXCLUDED_FIELDS
    fields |= EXTRACTION_RESPONSE_META_FIELDS
    return sorted(fields)


def _schema_instruction(modality: str) -> str:
    """Build an exact-key schema instruction for the extractor."""
    fields = _response_fields_for_modality(modality)
    if not fields:
        return ""

    field_list = ", ".join(f"`{field}`" for field in fields)
    return (
        "Use only these exact JSON field names when returning extracted data. "
        "Do not invent alternative key names, synonyms, or nested structures. "
        "If a field is absent or unclear, leave it null or omit it. "
        f"Allowed keys: {field_list}."
    )


def _has_meaningful_value(value: Any) -> bool:
    """Return True when a parsed extraction value should count as populated."""
    if value is None:
        return False
    if isinstance(value, str):
        return bool(value.strip())
    if isinstance(value, list):
        return any(_has_meaningful_value(item) for item in value)
    if isinstance(value, dict):
        return any(_has_meaningful_value(item) for item in value.values())
    return True


def _to_float(value: Any) -> float | None:
    """Best-effort numeric coercion for extracted values."""
    if value is None:
        return None
    if isinstance(value, (int, float)):
        if math.isfinite(value):
            return float(value)
        return None
    if isinstance(value, str):
        match = re.search(r"-?\d+(?:\.\d+)?", value.replace(",", ""))
        if match:
            return float(match.group(0))
    return None


def _alias_token(key: str) -> str:
    """Normalize an extracted key/label for alias matching."""
    return re.sub(r"[^a-z0-9]+", "_", key.strip().lower()).strip("_")


def _normalize_numeric_value(key: str, value: Any) -> float | None:
    """Coerce unit-bearing extracted numeric values into floats."""
    number = _to_float(value)
    if number is None:
        return None
    if key == "height" and 0 < number < 3:
        # Height is stored in centimetres; reports sometimes return metres.
        return number * 100
    return number


def _unit_hint(value: Any, explicit_unit: str | None = None) -> str | None:
    """Detect the explicit unit attached to an extracted measurement."""
    if explicit_unit:
        return explicit_unit.lower()
    if not isinstance(value, str):
        return None
    text = value.lower().replace("²", "2")
    if re.search(r"\bcm\s*/\s*s\b|\bcm/s\b|\bcm\s*s-1\b|\bcm\s*s\^-1\b", text):
        return "cm/s"
    if re.search(r"(?<!c)\bm\s*/\s*s\b|(?<!c)\bm/s\b|\bm\s*s-1\b|\bm\s*s\^-1\b", text):
        return "m/s"
    if re.search(r"\bmm\s*hg\b|\bmmhg\b", text):
        return "mmhg"
    if re.search(r"\bms\b|\bmsec\b|millisecond", text):
        return "ms"
    if re.search(r"\bmm\b|millimet", text):
        return "mm"
    if re.search(r"\bcm\b|centimet", text):
        return "cm"
    if re.search(r"\bm\b|\bmetres?\b|\bmeters?\b", text):
        return "m"
    if re.search(r"\bs\b|\bsec\b|second", text):
        return "s"
    return None


def _normalise_echo_numeric_measurement(
    key: str,
    value: Any,
    explicit_unit: str | None = None,
) -> float | None:
    """Convert Echo measurements into the units used by the UI/schema."""
    number = _to_float(value)
    if number is None:
        return None

    unit = _unit_hint(value, explicit_unit)
    if key == "height":
        if unit == "m" or 0 < number < 3:
            return number * 100
        return number

    if key in ECHO_MM_FIELDS:
        if unit == "cm":
            return number * 10
        if unit == "m":
            return number * 1000
        if unit is None:
            if key in ECHO_SMALL_MOTION_MM_FIELDS and 0 < number < 4:
                return number * 10
            if key in ECHO_WALL_THICKNESS_MM_FIELDS and 0 < number < 2.5:
                return number * 10
            if (
                key not in ECHO_SMALL_MOTION_MM_FIELDS
                and key not in ECHO_WALL_THICKNESS_MM_FIELDS
                and 0 < number < 10
            ):
                return number * 10
        return number

    if key in ECHO_CM_PER_SEC_FIELDS:
        if unit == "m/s":
            return number * 100
        return number

    if key in ECHO_M_PER_SEC_FIELDS:
        if unit == "cm/s":
            return number / 100
        return number

    if key in ECHO_CM_FIELDS:
        if unit == "mm":
            return number / 10
        if unit == "m":
            return number * 100
        return number

    if key in ECHO_MS_FIELDS:
        if unit == "s":
            return number * 1000
        return number

    return number


def _apply_aliases(payload: dict[str, Any], alias_map: dict[str, str]) -> None:
    """Copy common alias keys onto canonical schema keys when missing."""
    for source_key, source_value in list(payload.items()):
        target_key = alias_map.get(source_key) or alias_map.get(_alias_token(source_key))
        if not target_key:
            continue
        _set_if_missing(payload, target_key, source_value)


def _apply_echo_unit_aliases(payload: dict[str, Any]) -> None:
    """Map Echo aliases that encode source units onto canonical fields."""
    for source_key, source_value in list(payload.items()):
        target = ECHO_UNIT_ALIAS_MAP.get(source_key) or ECHO_UNIT_ALIAS_MAP.get(_alias_token(source_key))
        if not target:
            continue
        target_key, source_unit = target
        converted = _normalise_echo_numeric_measurement(target_key, source_value, source_unit)
        if converted is not None:
            _set_if_missing(payload, target_key, converted)


def _normalize_rhc_payload(payload: dict[str, Any]) -> None:
    """Normalize RHC extraction aliases and numeric values for review/save."""
    _apply_aliases(payload, RHC_ALIAS_MAP)

    normalized_date = _normalize_date_string(payload.get("date_rhc"))
    if normalized_date:
        payload["date_rhc"] = normalized_date

    for key in RHC_NUMERIC_FIELDS:
        value = payload.get(key)
        if not _has_meaningful_value(value):
            continue
        number = _normalize_numeric_value(key, value)
        if number is not None:
            payload[key] = number


def _normalize_echo_text_value(value: Any) -> str | None:
    """Convert Echo list/dict helper fields into readable review text."""
    if value is None:
        return None
    if isinstance(value, list):
        parts = [str(item).strip() for item in value if str(item).strip()]
        return "\n".join(parts) if parts else None
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    text = str(value).strip()
    return text or None


def _normalize_echo_payload(payload: dict[str, Any]) -> None:
    """Normalize Echo extraction aliases onto the review/save schema."""
    _apply_echo_unit_aliases(payload)
    _apply_aliases(payload, ECHO_ALIAS_MAP)

    for key in ECHO_NUMERIC_FIELDS:
        value = payload.get(key)
        if not _has_meaningful_value(value):
            continue
        number = _normalise_echo_numeric_measurement(key, value)
        if number is not None:
            payload[key] = number

    for text_key in ("conclusion", "conc_items", "narrative", "uncertain", "ai_warnings", "ai_raw_text"):
        text_value = _normalize_echo_text_value(payload.get(text_key))
        if text_value is not None:
            payload[text_key] = text_value


def _measurement_lines(text: str) -> list[str]:
    """Return normalized source lines for deterministic report backfills."""
    return [
        _normalize_whitespace(line.replace("\t", " ").replace("|", " | "))
        for line in text.splitlines()
        if _normalize_whitespace(line)
    ]


def _pipe_cells(line: str, *, keep_empty: bool = False) -> list[str]:
    """Split an antiword/docx table row into normalized pipe-delimited cells."""
    cells = [_normalize_whitespace(cell) for cell in line.strip().strip("|").split("|")]
    return cells if keep_empty else [cell for cell in cells if cell]


def _pipe_label_matches(cell: str, label_tokens: set[str], *, exact: bool = False) -> bool:
    token = _alias_token(cell.rstrip(":"))
    if not token:
        return False
    if exact:
        return token in label_tokens
    return any(token == label or token.startswith(f"{label}_") for label in label_tokens)


def _pipe_value_after_labels(
    text: str,
    labels: list[str],
    *,
    exact: bool = False,
) -> tuple[bool, str | None]:
    """Return the patient value immediately after a pipe-table label.

    The boolean records whether the label was found. This matters because a
    blank measured cell followed by a reference range must not fall through to
    the looser regex parser and be misread as a patient value.
    """
    label_tokens = {_alias_token(label) for label in labels}
    for line in text.splitlines():
        if "|" not in line:
            continue
        cells = _pipe_cells(line, keep_empty=True)
        for idx, cell in enumerate(cells):
            if not _pipe_label_matches(cell, label_tokens, exact=exact):
                continue
            if idx + 1 >= len(cells):
                return True, None
            next_cell = cells[idx + 1]
            if not next_cell:
                return True, None
            return True, next_cell.strip(" :")
    return False, None


def _set_echo_text_from_pipe(
    payload: dict[str, Any],
    key: str,
    text: str,
    labels: list[str],
    *,
    exact: bool = False,
    overwrite: bool = False,
) -> None:
    """Backfill a short Echo text/header field from a pipe-table row."""
    if _has_meaningful_value(payload.get(key)) and not overwrite:
        return
    found, value = _pipe_value_after_labels(text, labels, exact=exact)
    if found and value:
        _set_if_missing(payload, key, value, overwrite=overwrite)


def _extract_echo_rhythm(text: str) -> str | None:
    """Recover rhythm when the antiword table splits `Rhythm during study`."""
    lines = text.splitlines()
    for idx, line in enumerate(lines):
        if "rhythm during" not in line.lower():
            continue
        cells = _pipe_cells(line)
        rhythm_parts: list[str] = []
        if len(cells) > 1:
            rhythm_parts.append(cells[1])
        if idx + 1 < len(lines) and "study" in lines[idx + 1].lower():
            next_cells = _pipe_cells(lines[idx + 1])
            if len(next_cells) > 1:
                rhythm_parts.append(next_cells[1])
        rhythm = _normalize_whitespace(" ".join(rhythm_parts))
        if rhythm:
            return rhythm
    return None


def _extract_echo_image_quality(text: str) -> str | None:
    """Recover multi-line image-quality descriptions from converted tables."""
    lines = text.splitlines()
    for idx, line in enumerate(lines):
        if "image quality" not in line.lower():
            continue
        cells = _pipe_cells(line)
        parts: list[str] = []
        for cell_idx, cell in enumerate(cells):
            if _alias_token(cell) == "image_quality" and cell_idx + 1 < len(cells):
                parts.append(cells[cell_idx + 1])
                break
        for extra_line in lines[idx + 1 : idx + 4]:
            if re.search(r"\b(verbal consent|conclusions|height|bp)\b", extra_line, flags=re.IGNORECASE):
                break
            extra_cells = _pipe_cells(extra_line)
            if extra_cells:
                tail = extra_cells[-1]
                if not re.search(r"^(study|rhythm)$", tail, flags=re.IGNORECASE):
                    parts.append(tail)
        value = _normalize_whitespace(" ".join(parts))
        if value:
            return value
    return None


def _extract_echo_conclusions(text: str) -> str | None:
    """Extract and unwrap the visible Echo conclusions block."""
    lines = text.splitlines()
    collecting = False
    items: list[str] = []
    stop_pattern = re.compile(
        r"\b(normal .*dimensions|lvidd|ivsd|lvedv|lv mass|rwt|la max|la vol|ra area|rvd1|rvd2|rvd3|rvot2)\b",
        flags=re.IGNORECASE,
    )
    for line in lines:
        clean = " ".join(_pipe_cells(line)) if "|" in line else _normalize_whitespace(line)
        if not clean:
            continue
        if not collecting:
            if clean.lower() == "conclusions":
                collecting = True
            continue
        if stop_pattern.search(clean):
            break
        if items and not re.search(r"[.!?]$", items[-1]):
            items[-1] = _normalize_whitespace(f"{items[-1]} {clean}")
        else:
            items.append(clean)
    return "\n".join(items).strip() or None


def _pipe_value_before_continuation(
    text: str,
    labels: list[str],
    continuation_pattern: str,
) -> str | None:
    """Read values where a table label wraps onto the next converted line."""
    lines = text.splitlines()
    for idx, line in enumerate(lines[:-1]):
        if not re.search(continuation_pattern, lines[idx + 1], flags=re.IGNORECASE):
            continue
        found, value = _pipe_value_after_labels(line, labels)
        if found and value:
            return value
    return None


def _label_regex(label: str) -> str:
    """Build a tolerant regex for a measurement label."""
    escaped = re.escape(label.strip())
    escaped = escaped.replace(r"\ ", r"\s+")
    return escaped


def _first_fragment_after_labels(text: str, labels: list[str]) -> str | None:
    """Find the text fragment after the first matching measurement label."""
    found, pipe_value = _pipe_value_after_labels(text, labels)
    if found:
        return pipe_value

    for line in _measurement_lines(text):
        for label in labels:
            pattern = rf"(?i)(?:^|[\s|:;,(]){_label_regex(label)}(?:\b|[\s|:;,)'])"
            match = re.search(pattern, line)
            if not match:
                continue
            fragment = line[match.end():].strip(" |:")
            if fragment:
                return fragment
    return None


def _first_number_after_labels(text: str, labels: list[str]) -> tuple[float, str] | None:
    """Extract the first patient-specific number following any label."""
    fragment = _first_fragment_after_labels(text, labels)
    if not fragment:
        return None
    # Drop leading unit/reference parentheticals before reading the measured value.
    cleaned = re.sub(r"^\s*(?:\([^)]*\)\s*)+", "", fragment)
    cleaned = re.sub(r"^\s*(?:value|measured|result)\s*[:=|-]?\s*", "", cleaned, flags=re.IGNORECASE)
    number = _to_float(cleaned)
    if number is None:
        return None
    return number, fragment


def _set_echo_numeric_from_text(
    payload: dict[str, Any],
    key: str,
    source_text: str,
    labels: list[str],
    *,
    overwrite: bool = False,
) -> None:
    """Backfill an Echo numeric field from source text."""
    if _has_meaningful_value(payload.get(key)) and not overwrite:
        return
    match = _first_number_after_labels(source_text, labels)
    if not match:
        return
    _number, fragment = match
    converted = _normalise_echo_numeric_measurement(key, fragment)
    if converted is not None:
        _set_if_missing(payload, key, converted, overwrite=overwrite)


def _find_echo_descriptor(text: str, patterns: list[tuple[str, str]]) -> str | None:
    """Return the first descriptor supported by a regex pattern."""
    for pattern, value in patterns:
        if re.search(pattern, text, flags=re.IGNORECASE):
            return value
    return None


def _echo_valve_severity(text: str, valve_terms: list[str], lesion_terms: list[str]) -> str | None:
    """Extract a concise valve lesion severity from narrative text."""
    valve = r"|".join(re.escape(term) for term in valve_terms)
    lesion = r"|".join(re.escape(term) for term in lesion_terms)
    severity = r"(trivial|trace|mild(?:\s+to\s+moderate)?|moderate(?:\s+to\s+severe)?|severe)"
    patterns = [
        rf"\b(?P<sev>{severity})\s+(?:\w+\s+){{0,3}}(?:{valve})\s+(?:{lesion})\b",
        rf"\b(?:{valve})\s+(?:{lesion})\s+(?:is\s+)?(?P<sev>{severity})\b",
        rf"\b(?P<sev>{severity})\s+(?:MR|MS|TR|TS|AR|AS|PR|PS)\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            raw = _normalize_whitespace(match.group("sev"))
            return raw[:1].upper() + raw[1:].lower()
    if re.search(rf"\bno\s+(?:significant\s+)?(?:{valve})\s+(?:{lesion})\b", text, flags=re.IGNORECASE):
        return "None"
    return None


def _derive_echo_classification(payload: dict[str, Any], source_text: str) -> None:
    """Backfill broad Echo classification fields from extracted findings."""
    combined = "\n".join(
        str(part)
        for part in [
            payload.get("conclusion"),
            payload.get("conc_items"),
            payload.get("narrative"),
            source_text,
        ]
        if _has_meaningful_value(part)
    )
    lowered = combined.lower()

    if not _has_meaningful_value(payload.get("ph_prob")):
        if re.search(r"\bhigh\s+probability\b.*pulmonary hypertension|pulmonary hypertension.*\bhigh\s+probability\b", lowered):
            payload["ph_prob"] = "High"
        elif re.search(r"\bintermediate\s+probability\b.*pulmonary hypertension|pulmonary hypertension.*\bintermediate\s+probability\b", lowered):
            payload["ph_prob"] = "Intermediate"
        elif re.search(r"\blow\s+probability\b.*pulmonary hypertension|pulmonary hypertension.*\blow\s+probability\b", lowered):
            payload["ph_prob"] = "Low"

    if not _has_meaningful_value(payload.get("case_type")):
        normal_chambers = all(
            str(payload.get(key) or "").lower() == "normal"
            for key in ("lv_size", "rv_size", "la_size", "ra_size")
        )
        no_significant_valves = all(
            str(payload.get(key) or "").lower() == "none"
            for key in ("mr_grade", "tr_grade", "as_grade", "ar_grade")
        )
        if "pulmonary hypertension" in lowered or str(payload.get("ph_prob") or "").lower() in {"high", "intermediate"}:
            payload["case_type"] = "Pulmonary hypertension / right heart"
        elif any(term in lowered for term in ["severe aortic stenosis", "severe mitral", "moderate mitral", "moderate aortic", "valve"]):
            payload["case_type"] = "Valve"
        elif any(term in lowered for term in ["cardiomyopathy", "heart failure", "systolic dysfunction", "impaired lv"]):
            payload["case_type"] = "Heart failure / cardiomyopathy"
        elif "normal echocardiogram" in lowered or "normal echo" in lowered or (normal_chambers and no_significant_valves):
            payload["case_type"] = "Normal"

    if not _has_meaningful_value(payload.get("primary_dx")):
        if "pulmonary hypertension" in lowered:
            payload["primary_dx"] = "Pulmonary hypertension"
        elif payload.get("mr_grade") in {"Moderate", "Moderate to severe", "Severe"}:
            payload["primary_dx"] = f"{payload['mr_grade']} MR"
        elif payload.get("as_grade") in {"Moderate", "Moderate to severe", "Severe"}:
            payload["primary_dx"] = f"{payload['as_grade']} AS"
        elif payload.get("lvef") is not None and _to_float(payload.get("lvef")) is not None:
            lvef = _to_float(payload.get("lvef"))
            if lvef is not None and lvef < 50:
                payload["primary_dx"] = "LV systolic dysfunction"
        elif str(payload.get("case_type") or "") == "Normal":
            payload["primary_dx"] = "Normal echo"

    if not _has_meaningful_value(payload.get("secondary_path")):
        if re.search(r"\bright ventricular (?:dilatation|dilation|dysfunction|impairment)\b", lowered):
            payload["secondary_path"] = "Right ventricular dilatation and dysfunction"
        elif re.search(r"\bascending aort(?:a|ic) (?:dilatation|dilation)\b", lowered):
            payload["secondary_path"] = "Ascending aortic dilatation"


def _backfill_echo_from_text(payload: dict[str, Any], source_text: str | None) -> None:
    """Apply deterministic Echo backfills from converted report text."""
    text = source_text or ""
    if not text:
        return

    for date_key, labels in [
        ("study_date", ["study date", "echo date"]),
        ("report_date", ["report date", "reported date"]),
    ]:
        if not _has_meaningful_value(payload.get(date_key)):
            fragment = _first_fragment_after_labels(text, labels)
            normalized_date = _normalize_date_string(fragment)
            if normalized_date:
                payload[date_key] = normalized_date
    if not _has_meaningful_value(payload.get("study_date")):
        found, fragment = _pipe_value_after_labels(text, ["date"], exact=True)
        normalized_date = _normalize_date_string(fragment if found else None)
        if normalized_date:
            payload["study_date"] = normalized_date

    for key, labels in {
        "consultant": ["consultant"],
        "ward_op": ["ward/op", "ward op"],
        "study_reason": ["purpose of study", "study reason", "indication"],
        "reported_by": ["reported by", "operator"],
    }.items():
        _set_echo_text_from_pipe(payload, key, text, labels)

    if not _has_meaningful_value(payload.get("rhythm")):
        rhythm = _extract_echo_rhythm(text)
        _set_if_missing(payload, "rhythm", rhythm)
    if not _has_meaningful_value(payload.get("image_quality")):
        image_quality = _extract_echo_image_quality(text)
        _set_if_missing(payload, "image_quality", image_quality)

    conclusions = _extract_echo_conclusions(text)
    if conclusions:
        _set_if_missing(payload, "conclusion", conclusions)
        _set_if_missing(payload, "conc_items", conclusions)

    hr_range = re.search(
        r"\b(?:heart rate|hr|rate)\b[^\d]{0,20}(\d+(?:\.\d+)?)\s*(?:-|to|\u2013|\u2014)\s*(\d+(?:\.\d+)?)\s*(?:bpm)?",
        text,
        flags=re.IGNORECASE,
    )
    if hr_range:
        mean_hr = (float(hr_range.group(1)) + float(hr_range.group(2))) / 2
        _set_if_missing(payload, "hr", mean_hr, overwrite=True)
    else:
        hr_single = re.search(
            r"\b(?:rhythm\s*[:|]?\s*rate|heart rate|hr|rate)\b[^\d]{0,40}(\d+(?:\.\d+)?)\s*(?:bpm)?",
            text,
            flags=re.IGNORECASE,
        )
        if hr_single:
            _set_if_missing(payload, "hr", float(hr_single.group(1)))

    echo_numeric_specs: dict[str, list[str]] = {
        "hr": ["heart rate", "hr", "rhythm rate"],
        "height": ["height", "height cm"],
        "weight": ["weight", "weight kg"],
        "bsa": ["bsa", "body surface area"],
        "lvef": ["lvef", "lv ef", "ejection fraction"],
        "gls": ["gls", "global longitudinal strain"],
        "mapse": ["mapse"],
        "med_s": ["medial s'", "med s'", "septal s'"],
        "lat_s": ["lateral s'", "lat s'"],
        "sept_e": ["septal e'", "sept e'"],
        "lat_e": ["lateral e'", "lat e'"],
        "avg_e_ep": ["average e/e'", "avg e/e'", "mean e/e'"],
        "sept_e_ep": ["septal e/e'", "sept e/e'"],
        "mv_e": ["mv e", "mitral e", "e wave"],
        "mv_a": ["mv a", "mitral a", "a wave"],
        "e_a": ["e/a", "e:a"],
        "dt_ms": ["deceleration time", "dt"],
        "tapse": ["tapse"],
        "rv_s": ["rv s'", "rv s prime", "tricuspid annular s'", "tricuspid tdi"],
        "fac": ["fac", "fractional area change"],
        "tr_vmax": ["tr vmax", "tr v max", "tr velocity", "tr peak velocity"],
        "rvsp": ["rvsp", "estimated rvsp", "right ventricular systolic pressure"],
        "rap": ["rap", "estimated rap", "right atrial pressure"],
        "av_vmax": ["av vmax", "aortic valve vmax", "aortic vmax"],
        "av_pk_grad": ["av peak gradient", "aortic valve peak gradient", "peak gradient"],
        "av_mn_grad": ["av mean gradient", "aortic valve mean gradient", "mean gradient"],
        "ava": ["ava", "aortic valve area"],
        "ar_pht": ["ar pht", "ar pressure half-time", "pressure half time"],
        "pat": ["pat", "pvat", "pulmonary acceleration time"],
        "pv_vmax": ["pv vmax", "pulmonary valve vmax"],
        "lvidd": ["lvidd", "lvid d", "lv internal diameter diastole"],
        "lvids": ["lvids", "lvid s", "lv internal diameter systole"],
        "ivsd": ["ivsd", "ivs diastole", "septal wall thickness"],
        "lvpwd": ["lvpwd", "posterior wall thickness"],
        "lvedvi": ["lvedvi", "lv edvi", "lv edv index", "lvedv indexed", "lv edv indexed"],
        "lvesvi": ["lvesvi", "lv esvi", "lv esv index", "lvesv indexed", "lv esv indexed"],
        "rwt": ["rwt", "relative wall thickness"],
        "lvmi": ["lvmi", "lv mass index", "lv mass indexed"],
        "la_diam": ["la diameter", "left atrial diameter"],
        "la_vol": ["la volume", "left atrial volume"],
        "la_voli": ["lavi", "la volume index", "left atrial volume index", "la vol indexed"],
        "ra_area": ["ra area", "right atrial area"],
        "ra_areai": ["ra area index", "ra area indexed", "right atrial area index"],
        "rvd1": ["rvd1", "rv basal diameter", "rv base"],
        "rvd2": ["rvd2", "rv mid diameter", "rv mid"],
        "rvd3": ["rvd3", "rv longitudinal diameter", "rv length"],
        "rvot2": ["rvot", "rvot2"],
        "ao_ann": ["aortic annulus", "ao annulus"],
        "ao_sinus": ["sinus of valsalva", "aortic sinus", "ao sinus", "ao @ sinus"],
        "stj_mm": ["sinotubular junction", "stj", "sino tubular", "sino tubular junct"],
        "asc_ao_prox": ["ascending aorta proximal", "proximal ascending aorta", "proximal ascending ao"],
        "asc_ao_mid": ["ascending aorta mid", "mid ascending aorta", "right para mid asc ao"],
        "ao_arch": ["aortic arch", "ao arch"],
        "main_pa_mm": ["main pulmonary artery", "mpa diameter", "main pa"],
        "lvot_diam": ["lvot diameter", "lvot diam"],
        "lvot_vel": ["lvot velocity", "lvot vel"],
        "lvot_vti": ["lvot vti"],
        "dvi": ["dvi", "dimensionless index"],
        "av_vti": ["av vti", "aortic valve vti"],
    }
    for key, labels in echo_numeric_specs.items():
        _set_echo_numeric_from_text(payload, key, text, labels, overwrite=key == "hr" and bool(hr_range))

    if not _has_meaningful_value(payload.get("asc_ao_prox")):
        proximal_ao = _pipe_value_before_continuation(text, ["proximal"], r"\bascending\s+ao\b")
        converted = _normalise_echo_numeric_measurement("asc_ao_prox", proximal_ao)
        if converted is not None:
            payload["asc_ao_prox"] = converted

    bsa = _compute_bsa_from_payload(payload)
    if bsa is not None:
        _set_if_missing(payload, "bsa", round(bsa, 2))

    descriptor_text = "\n".join([str(payload.get("narrative") or ""), str(payload.get("conclusion") or ""), text])
    descriptor_specs: dict[str, list[tuple[str, str]]] = {
        "lv_size": [
            (r"\bnormal size left ventricle\b|\bnormal size lv\b", "Normal"),
            (r"\bleft ventricle\b.{0,80}\b(?:normal in size|normal size|not dilated|non-dilated)\b|\blv\b.{0,50}\b(?:normal in size|normal size|not dilated|non-dilated)\b", "Normal"),
            (r"\bleft ventricle\b.{0,80}\bmild(?:ly)? dilat(?:ed|ation)\b", "Mildly dilated"),
            (r"\bleft ventricle\b.{0,80}\bmoderate(?:ly)? dilat(?:ed|ation)\b", "Moderately dilated"),
            (r"\bleft ventricle\b.{0,80}\bsevere(?:ly)? dilat(?:ed|ation)\b", "Severely dilated"),
        ],
        "lv_wall": [
            (r"\bnormal wall thickness\b|\bwall thickness\b.{0,20}\bnormal\b", "Normal"),
            (r"\bmild(?:ly)? (?:concentric |eccentric )?(?:lv )?hypertroph(?:y|ied)\b", "Mild hypertrophy"),
            (r"\bmoderate(?:ly)? (?:concentric |eccentric )?(?:lv )?hypertroph(?:y|ied)\b", "Moderate hypertrophy"),
            (r"\bsevere(?:ly)? (?:concentric |eccentric )?(?:lv )?hypertroph(?:y|ied)\b", "Severe hypertrophy"),
        ],
        "rv_size": [
            (r"\bnormal size right ventricle\b|\bnormal size rv\b", "Normal"),
            (r"\bright ventricle\b.{0,80}\b(?:normal in size|normal size|not dilated|non-dilated)\b|\brv\b.{0,50}\b(?:normal in size|normal size|not dilated|non-dilated)\b", "Normal"),
            (r"\bright ventric(?:le|ular)\b.{0,80}\bmild(?:ly)? dilat(?:ed|ation)\b", "Mildly dilated"),
            (r"\bright ventric(?:le|ular)\b.{0,80}\bmoderate(?:ly)? dilat(?:ed|ation)\b", "Moderately dilated"),
            (r"\bright ventric(?:le|ular)\b.{0,80}\bsevere(?:ly)? dilat(?:ed|ation)\b", "Severely dilated"),
            (r"\bright ventric(?:le|ular)\b.{0,80}\bdilat(?:ed|ation)\b", "Dilated"),
        ],
        "lv_fn": [
            (r"\bleft ventricle\b.{0,120}\bnormal systolic function\b|\bnormal systolic function\b.{0,40}\bleft ventricle\b", "Normal"),
            (r"\b(?:normal|preserved|good)\s+lv systolic function\b|\blv\b.{0,60}\b(?:systolic )?function\b.{0,40}\b(?:normal|preserved|good)\b|\bleft ventricular systolic function\b.{0,40}\b(?:normal|preserved|good)\b", "Normal"),
            (r"\blv\b.{0,60}\b(?:mild(?:ly)? impaired|mild dysfunction)\b", "Mildly impaired"),
            (r"\blv\b.{0,60}\b(?:moderate(?:ly)? impaired|moderate dysfunction)\b", "Moderately impaired"),
            (r"\blv\b.{0,60}\b(?:severe(?:ly)? impaired|severe dysfunction)\b", "Severely impaired"),
        ],
        "rv_fn": [
            (r"\bright ventricle\b.{0,120}\bnormal systolic function\b|\bnormal systolic function\b.{0,40}\bright ventricle\b", "Normal"),
            (r"\brv\b.{0,60}\b(?:systolic )?function\b.{0,40}\b(?:normal|preserved|good)\b|\bright ventricular systolic function\b.{0,40}\b(?:normal|preserved|good)\b", "Normal"),
            (r"\brv\b.{0,60}\b(?:mild(?:ly)? impaired|mild dysfunction|reduced)\b", "Mildly impaired"),
            (r"\brv\b.{0,60}\b(?:moderate(?:ly)? impaired|moderate dysfunction)\b", "Moderately impaired"),
            (r"\brv\b.{0,60}\b(?:severe(?:ly)? impaired|severe dysfunction)\b", "Severely impaired"),
        ],
        "la_size": [
            (r"\bnormal size atria\b", "Normal"),
            (r"\bleft atri(?:um|a)\b.{0,60}\bnormal\b.{0,20}\bsize\b|\bla\b.{0,30}\bnormal size\b", "Normal"),
            (r"\bleft atri(?:um|a)\b.{0,60}\bmild(?:ly)? dilat(?:ed|ation)\b", "Mildly dilated"),
            (r"\bleft atri(?:um|a)\b.{0,60}\bmoderate(?:ly)? dilat(?:ed|ation)\b", "Moderately dilated"),
            (r"\bleft atri(?:um|a)\b.{0,60}\bsevere(?:ly)? dilat(?:ed|ation)\b", "Severely dilated"),
        ],
        "ra_size": [
            (r"\bnormal size atria\b", "Normal"),
            (r"\bright atri(?:um|a)\b.{0,60}\bnormal\b.{0,20}\bsize\b|\bra\b.{0,30}\bnormal size\b", "Normal"),
            (r"\bright atri(?:um|a)\b.{0,60}\bmild(?:ly)? dilat(?:ed|ation)\b", "Mildly dilated"),
            (r"\bright atri(?:um|a)\b.{0,60}\bmoderate(?:ly)? dilat(?:ed|ation)\b", "Moderately dilated"),
            (r"\bright atri(?:um|a)\b.{0,60}\bsevere(?:ly)? dilat(?:ed|ation)\b", "Severely dilated"),
        ],
    }
    for key, patterns in descriptor_specs.items():
        _set_if_missing(payload, key, _find_echo_descriptor(descriptor_text, patterns))

    if re.search(r"\bno significant valvular abnormalities\b", descriptor_text, flags=re.IGNORECASE):
        for key in ("ms_grade", "mr_grade", "ts_grade", "tr_grade", "as_grade", "ar_grade", "ps_grade", "pr_grade"):
            _set_if_missing(payload, key, "None")
        for key in ("mv_desc", "tv_desc", "av_desc", "pv_desc"):
            _set_if_missing(payload, key, "Normal")

    if not _has_meaningful_value(payload.get("ivc_size")):
        found, ivc_value = _pipe_value_after_labels(text, ["ivc"], exact=False)
        if found and ivc_value:
            ivc_number = _to_float(ivc_value)
            if ivc_number is not None:
                payload["ivc_size"] = _format_number(ivc_number, 1)

    for key, valve_terms, lesion_terms in [
        ("mr_grade", ["mitral", "MR"], ["regurgitation", "regurgitant", "MR"]),
        ("ms_grade", ["mitral", "MS"], ["stenosis", "MS"]),
        ("tr_grade", ["tricuspid", "TR"], ["regurgitation", "regurgitant", "TR"]),
        ("ts_grade", ["tricuspid", "TS"], ["stenosis", "TS"]),
        ("ar_grade", ["aortic", "AR"], ["regurgitation", "regurgitant", "AR"]),
        ("as_grade", ["aortic", "AS"], ["stenosis", "AS"]),
        ("pr_grade", ["pulmonary", "pulmonic", "PR"], ["regurgitation", "regurgitant", "PR"]),
        ("ps_grade", ["pulmonary", "pulmonic", "PS"], ["stenosis", "PS"]),
    ]:
        _set_if_missing(payload, key, _echo_valve_severity(descriptor_text, valve_terms, lesion_terms))

    bool_specs = {
        "sept_flat": r"septal flattening|flattened septum",
        "sept_bounce": r"septal bounce",
        "d_shaped_lv": r"d[- ]shaped lv|d shaped left ventricle",
        "pa_dilated": r"pulmonary artery (?:is )?(?:dilated|dilatation)|dilated main pulmonary artery",
        "peric_eff": r"pericardial effusion",
        "shunt": r"\bshunt\b|colour flow across the interatrial septum",
    }
    for key, pattern in bool_specs.items():
        if _has_meaningful_value(payload.get(key)):
            continue
        if re.search(rf"\b(?:no|without|absent)\b.{0,20}(?:{pattern})", descriptor_text, flags=re.IGNORECASE):
            payload[key] = 0
        elif re.search(pattern, descriptor_text, flags=re.IGNORECASE):
            payload[key] = 1

    if not _has_meaningful_value(payload.get("ias_intact")):
        if re.search(r"interatrial septum.{0,40}(?:intact|no colour flow|no color flow)", descriptor_text, flags=re.IGNORECASE):
            payload["ias_intact"] = 1
        elif re.search(r"interatrial septum.{0,40}(?:defect|shunt|pfo|asd)", descriptor_text, flags=re.IGNORECASE):
            payload["ias_intact"] = 0

    _derive_echo_classification(payload, text)


def _format_number(value: float, decimals: int = 1) -> str:
    """Format derived numbers without unnecessary trailing zeros."""
    rounded = round(value, decimals)
    if abs(rounded - round(rounded)) < 1e-9:
        return str(int(round(rounded)))
    return f"{rounded:.{decimals}f}".rstrip("0").rstrip(".")


def _normalize_whitespace(value: str) -> str:
    """Collapse repeated whitespace for short text fields."""
    return re.sub(r"\s+", " ", value).strip()


def _dedupe_preserve_order(values: list[str]) -> list[str]:
    """Remove duplicate extracted DOCX lines while preserving report order."""
    seen: set[str] = set()
    result: list[str] = []
    for value in values:
        line = _normalize_whitespace(value)
        if not line or line in seen:
            continue
        seen.add(line)
        result.append(line)
    return result


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract visible paragraph and table-cell text from a DOCX file."""
    from io import BytesIO

    from docx import Document

    doc = Document(BytesIO(file_bytes))
    lines: list[str] = []

    def add_paragraphs(paragraphs: Any) -> None:
        for paragraph in paragraphs:
            text = _normalize_whitespace(paragraph.text)
            if text:
                lines.append(text)

    def add_table(table: Any) -> None:
        for row in table.rows:
            row_values: list[str] = []
            for cell in row.cells:
                cell_parts: list[str] = []
                for paragraph in cell.paragraphs:
                    text = _normalize_whitespace(paragraph.text)
                    if text:
                        cell_parts.append(text)
                for nested_table in cell.tables:
                    add_table(nested_table)
                cell_text = " ".join(cell_parts).strip()
                if cell_text:
                    row_values.append(cell_text)
            if row_values:
                lines.append(" | ".join(row_values))

    add_paragraphs(doc.paragraphs)
    for table in doc.tables:
        add_table(table)
    for section in doc.sections:
        add_paragraphs(section.header.paragraphs)
        for table in section.header.tables:
            add_table(table)
        add_paragraphs(section.footer.paragraphs)
        for table in section.footer.tables:
            add_table(table)

    return "\n".join(_dedupe_preserve_order(lines))


def _clean_document_text(text: str) -> str:
    """Normalize converted document text while preserving line breaks."""
    return "\n".join(
        line.strip()
        for line in text.splitlines()
        if line.strip()
    ).strip()


def _extract_doc_text(file_bytes: bytes) -> str:
    """Extract readable text from legacy Word .doc files."""
    if file_bytes.startswith(b"PK"):
        try:
            text = _extract_docx_text(file_bytes)
            if text.strip():
                return text
        except Exception as exc:  # pragma: no cover - defensive renamed-file fallback
            logger.info("DOC upload was not readable as DOCX: %s", exc)

    with tempfile.TemporaryDirectory() as temp_dir:
        doc_path = Path(temp_dir) / "upload.doc"
        doc_path.write_bytes(file_bytes)

        for tool_name in ("antiword", "catdoc"):
            tool_path = shutil.which(tool_name)
            if not tool_path:
                continue
            try:
                result = subprocess.run(
                    [tool_path, str(doc_path)],
                    capture_output=True,
                    text=True,
                    timeout=30,
                    check=False,
                )
            except (OSError, subprocess.TimeoutExpired) as exc:
                logger.warning("%s failed while reading DOC upload: %s", tool_name, exc)
                continue
            text = _clean_document_text(result.stdout or "")
            if result.returncode == 0 and text:
                return text
            logger.warning(
                "%s could not read DOC upload: returncode=%s stderr=%s",
                tool_name,
                result.returncode,
                (result.stderr or "")[:200],
            )

        office_path = shutil.which("soffice") or shutil.which("libreoffice")
        if office_path:
            try:
                result = subprocess.run(
                    [
                        office_path,
                        "--headless",
                        "--convert-to",
                        "txt:Text",
                        "--outdir",
                        temp_dir,
                        str(doc_path),
                    ],
                    capture_output=True,
                    text=True,
                    timeout=60,
                    check=False,
                )
                txt_path = Path(temp_dir) / "upload.txt"
                if result.returncode == 0 and txt_path.exists():
                    text = _clean_document_text(txt_path.read_text(encoding="utf-8", errors="ignore"))
                    if text:
                        return text
                logger.warning(
                    "LibreOffice could not read DOC upload: returncode=%s stderr=%s",
                    result.returncode,
                    (result.stderr or "")[:200],
                )
            except (OSError, subprocess.TimeoutExpired) as exc:
                logger.warning("LibreOffice failed while reading DOC upload: %s", exc)

    raise ExtractionParseError(
        "Could not read the legacy .doc file. Please save it as .docx or PDF and try again."
    )


def _normalize_date_string(value: Any) -> str | None:
    """Normalize common extracted date formats to dd/mm/yyyy."""
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None

    candidate = text
    if "T" in candidate:
        candidate = candidate.split("T", 1)[0]
    if " " in candidate and re.match(r"^\d{4}-\d{1,2}-\d{1,2}\b", candidate):
        candidate = candidate.split(" ", 1)[0]

    for fmt in (
        "%Y-%m-%d",
        "%d/%m/%Y",
        "%d-%m-%Y",
        "%d.%m.%Y",
        "%d %b %Y",
        "%d %B %Y",
        "%b %d, %Y",
        "%B %d, %Y",
    ):
        try:
            return datetime.strptime(candidate, fmt).strftime("%d/%m/%Y")
        except ValueError:
            continue

    match = re.search(r"\b(\d{4})-(\d{1,2})-(\d{1,2})\b", text)
    if match:
        return f"{int(match.group(3)):02d}/{int(match.group(2)):02d}/{int(match.group(1)):04d}"

    match = re.search(r"\b(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})\b", text)
    if match:
        day = int(match.group(1))
        month = int(match.group(2))
        year = int(match.group(3))
        if year < 100:
            year += 2000
        if 1 <= day <= 31 and 1 <= month <= 12:
            return f"{day:02d}/{month:02d}/{year:04d}"

    return None


def _set_if_missing(
    payload: dict[str, Any],
    key: str,
    value: Any,
    *,
    overwrite: bool = False,
) -> None:
    """Set a normalized payload field when it is currently blank."""
    if not _has_meaningful_value(value):
        return
    if overwrite or not _has_meaningful_value(payload.get(key)):
        payload[key] = value


def _compute_bsa_from_payload(payload: dict[str, Any]) -> float | None:
    """Compute Mosteller BSA from height/weight when both are available."""
    height_cm = _to_float(payload.get("height"))
    weight_kg = _to_float(payload.get("weight"))
    if height_cm and weight_kg and height_cm > 0 and weight_kg > 0:
        return math.sqrt((height_cm * weight_kg) / 3600)
    return None


def _derive_cmr_class(primary_dx: str | None, conclusions: str | None) -> str | None:
    """Backfill the broad CMR case class from the dominant diagnosis language."""
    text = " ".join(
        part.strip().lower()
        for part in [primary_dx or "", conclusions or ""]
        if part and part.strip()
    )
    if not text:
        return None
    supports_ph = bool(
        re.search(
            r"(?:cmr (?:features|findings)|features|findings).{0,80}support(?:s)?\s+(?:pulmonary hypertension|ph)\s+physiology",
            text,
        )
    )
    negates_ph_support = bool(
        re.search(
            r"\bno\s+cmr\s+(?:features|findings).{0,80}support(?:s)?\s+(?:pulmonary hypertension|ph)\s+physiology",
            text,
        )
    )
    if "pulmonary hypertension phenotype" in text or (supports_ph and not negates_ph_support):
        return "Pulmonary hypertension / right heart"
    ischaemic_text = re.sub(r"\bnon[- ]ischaem\w*\b|\bnon[- ]ischem\w*\b", " ", text)
    if any(
        token in ischaemic_text
        for token in [
            "infarct",
            "ischaemic scar",
            "ischemic scar",
            "ischaemic cardiomyopathy",
            "ischemic cardiomyopathy",
            "subendocardial",
            "transmural",
        ]
    ):
        return "Ischaemic"
    if any(token in text for token in ["pulmonary hypertension", "right heart"]):
        return "Pulmonary hypertension / right heart"
    if "normal" in text and "infarct" not in text and "scar" not in text:
        return "Normal"
    if any(token in text for token in ["myocarditis", "myopericarditis", "inflammatory"]):
        return "Myocarditis / inflammatory"
    if any(token in text for token in ["hypertrophic", "lvh"]):
        return "Hypertrophic cardiomyopathy / LVH"
    if any(token in text for token in ["amyloid", "sarcoid", "infiltr", "storage"]):
        return "Infiltrative / storage"
    if "pericard" in text:
        return "Pericardial disease"
    if "congenital" in text:
        return "Congenital heart disease"
    if any(token in text for token in ["regurg", "stenosis", "valve"]):
        return "Valve disease"
    if any(token in text for token in ["thrombus", "mass"]):
        return "Mass / thrombus"
    if any(token in text for token in ["prosthetic", "post-operative", "post operative", "surgery"]):
        return "Post-operative / prosthetic"
    if any(token in text for token in ["cardiomyopathy", "fibrosis"]):
        return "Non-ischaemic cardiomyopathy"
    return "Other"


def _extract_indication_from_text(text: str) -> str | None:
    """Extract a concise indication/clinical history string when explicitly present."""
    explicit_match = re.search(
        r"(?im)^\s*(?:indication|clinical indication|study reason|reason for (?:scan|study))\s*[:\-]\s*(.+)$",
        text,
    )
    if explicit_match:
        candidate = _normalize_whitespace(explicit_match.group(1))
        return candidate or None

    history_match = re.search(
        r"(?is)\bclinical history\b\s*:?\s*(.+?)(?:\n\s*\n|\n[A-Z][A-Za-z /&-]{1,40}:\s|\Z)",
        text,
    )
    if history_match:
        lines = [
            _normalize_whitespace(line)
            for line in history_match.group(1).splitlines()
            if _normalize_whitespace(line)
        ]
        if lines:
            return " ".join(lines[:2]).strip() or None

    return None


def _clean_text_for_ischaemic_inference(*parts: str) -> str:
    """Prepare narrative text for cautious ischaemic-scar inference."""
    text = " ".join(part for part in parts if part).lower()
    text = re.sub(
        r"rule out [^.:\n;]*?(?:infarct(?:ion)?|myocardial infarction|ischaem\w*|ischem\w*)[^.:\n;]*",
        " ",
        text,
    )
    return _normalize_whitespace(text)


def _has_positive_ischaemic_scar_evidence(
    payload: dict[str, Any],
    conclusions: str,
    source_text: str | None,
) -> bool:
    """Return True when the report supports a scar/infarct diagnosis."""
    evidence_parts = [
        conclusions,
        str(payload.get("lge") or ""),
        str(payload.get("fibrosis") or ""),
        str(payload.get("lge_pattern") or ""),
        str(payload.get("lge_location") or ""),
        str(payload.get("lge_transmurality") or ""),
        str(payload.get("fixed_defect") or ""),
        str(payload.get("perfusion_defect") or ""),
        str(payload.get("perfusion_coronary_territory") or ""),
        source_text or "",
    ]
    combined = _clean_text_for_ischaemic_inference(*evidence_parts)
    negative_markers = (
        "no evidence of infarct",
        "no evidence of any scar",
        "no scar",
        "no fibrosis",
        "no lge",
        "no late gadolinium enhancement",
        "no inducible ischaemia",
        "no inducible ischemia",
    )
    if any(marker in combined for marker in negative_markers):
        return False
    positive_markers = (
        "subendocardial",
        "transmural",
        "ischaemic scar",
        "ischemic scar",
        "prior myocardial infarction",
        "myocardial infarction",
        "infarct",
        "fixed defect",
    )
    return any(marker in combined for marker in positive_markers)


def _normalize_cmr_coded_values(payload: dict[str, Any], source_text: str | None) -> None:
    """Normalize free-text CMR outputs onto the house option vocabulary."""
    source_lower = (source_text or "").lower()

    contrast = str(payload.get("contrast") or "").strip()
    if (
        contrast
        and re.search(
            r"late gadolinium enhancement|\bgadolinium\b|\blge\b|^yes$|^present$|^true$",
            contrast,
            flags=re.IGNORECASE,
        )
        and re.search(r"late gadolinium enhancement|\bgadolinium\b|\blge\b", source_lower, flags=re.IGNORECASE)
    ):
        payload["contrast"] = "Gadolinium"

    flow = str(payload.get("flow") or "").strip()
    if flow and re.search(r"\b4d\s*[- ]?flow\b", flow, flags=re.IGNORECASE):
        payload["flow"] = "4D-flow"
    elif flow and re.search(r"\b2d\s*(?:[- ]?flow|-?pc)\b", flow, flags=re.IGNORECASE):
        payload["flow"] = "2D-flow"

    ph = str(payload.get("ph") or "").strip()
    if ph:
        lowered = ph.lower()
        if "low" in lowered or "unlikely" in lowered:
            payload["ph"] = "Low"
        elif "intermediate-high" in lowered or "intermediate high" in lowered:
            payload["ph"] = "Intermediate-high"
        elif "intermediate" in lowered:
            payload["ph"] = "Intermediate"
        elif "low-intermediate" in lowered or "low intermediate" in lowered or "mild" in lowered:
            payload["ph"] = "Low-intermediate"
        elif "high" in lowered:
            payload["ph"] = "High"

    primary_dx = str(payload.get("primary_dx") or "").strip()
    conclusions = str(payload.get("conclusions") or "").strip()
    supports_ph_physiology = bool(
        re.search(
            r"(?:cmr (?:features|findings)|features).{0,80}support(?:s)?\s+(?:pulmonary hypertension|ph)\s+physiology|support ph physiology",
            source_text or "",
            flags=re.IGNORECASE,
        )
    )
    has_ischaemic_scar_evidence = _has_positive_ischaemic_scar_evidence(
        payload,
        conclusions,
        source_text,
    )
    if not primary_dx and has_ischaemic_scar_evidence:
        payload["primary_dx"] = "Prior myocardial infarction / ischaemic scar"
    elif primary_dx:
        lowered = primary_dx.lower()
        if lowered in {"ischaemic", "ischemic"}:
            if has_ischaemic_scar_evidence:
                payload["primary_dx"] = "Prior myocardial infarction / ischaemic scar"
            else:
                payload["primary_dx"] = "Ischaemic cardiomyopathy"
        elif has_ischaemic_scar_evidence and any(
            token in lowered
            for token in ["non-ischaemic", "non-ischemic", "other", "normal", "cardiomyopathy"]
        ):
            payload["primary_dx"] = "Prior myocardial infarction / ischaemic scar"
        elif "pulmonary hypertension" in lowered and "phenotype" not in lowered:
            payload["primary_dx"] = "Pulmonary hypertension phenotype"

    primary_dx = str(payload.get("primary_dx") or "").strip()
    has_non_ischaemic_lge = bool(
        re.search(
            r"\bnon[- ]ischaemic\b.{0,80}\b(?:lge|fibrosis|scar)\b|\bmid[- ]wall lge\b",
            source_text or "",
            flags=re.IGNORECASE,
        )
    )
    if supports_ph_physiology and not has_ischaemic_scar_evidence:
        if primary_dx and primary_dx != "Pulmonary hypertension phenotype":
            _set_if_missing(payload, "secondary_dx", primary_dx)
        payload["primary_dx"] = "Pulmonary hypertension phenotype"
        if has_non_ischaemic_lge:
            _set_if_missing(payload, "secondary_dx", "Non-ischaemic myocardial fibrosis / scar")
    elif not primary_dx and has_non_ischaemic_lge:
        payload["primary_dx"] = "Non-ischaemic myocardial fibrosis / scar"


def _set_cmr_number_from_text(
    payload: dict[str, Any],
    key: str,
    source_text: str,
    labels: list[str],
    *,
    decimals: int = 1,
    overwrite: bool = False,
) -> None:
    """Backfill a CMR scalar measurement from source text as display text."""
    if _has_meaningful_value(payload.get(key)) and not overwrite:
        return
    match = _first_number_after_labels(source_text, labels)
    if not match:
        return
    number, _fragment = match
    _set_if_missing(payload, key, _format_number(number, decimals), overwrite=overwrite)


def _extract_cmr_date_from_text(text: str) -> str | None:
    """Recover the CMR study date from common report headers."""
    for pattern in [
        r"(?im)^\s*date of (?:examination|study|scan)\s*:?\s*(.+)$",
        r"(?im)^\s*study date\s*:?\s*(.+)$",
        r"(?im)^\s*radiology examination[^\n:]*:\s*(.+)$",
    ]:
        match = re.search(pattern, text)
        if not match:
            continue
        normalized = _normalize_date_string(match.group(1))
        if normalized:
            return normalized
    return None


def _extract_cmr_conclusions(text: str) -> str | None:
    """Extract the final CMR Conclusions section when OpenAI leaves it blank."""
    match = re.search(r"(?is)\bconclusions?\s*:?\s*(.+)$", text)
    if not match:
        return None
    chunk = match.group(1)
    stop = re.search(
        r"(?im)^\s*(?:reported by|report(?:ed)? by|signed by|electronically signed)\b",
        chunk,
    )
    if stop:
        chunk = chunk[: stop.start()]
    lines = [
        _normalize_whitespace(line)
        for line in chunk.splitlines()
        if _normalize_whitespace(line)
    ]
    return "\n".join(lines) if lines else None


def _backfill_cmr_2d_flow_table(payload: dict[str, Any], text: str) -> None:
    """Recover Aorta/Pulmonary paired values from standard `Flow (2D-PC)` tables."""
    if re.search(r"\bflow\s*\(\s*2d\s*-?\s*pc\s*\)", text, flags=re.IGNORECASE):
        _set_if_missing(payload, "flow", "2D-flow", overwrite=True)

    for line in _measurement_lines(text):
        lowered = line.lower()
        numbers = [float(n) for n in re.findall(r"-?\d+(?:\.\d+)?", line)]
        if len(numbers) < 2:
            continue
        if re.match(r"^forward flow\b", lowered):
            payload["ao_forward_volume"] = _format_number(numbers[0], 1)
            payload["pulmonary_forward_volume"] = _format_number(numbers[1], 1)
        elif re.match(r"^backward flow\b", lowered):
            payload["ao_backward_volume"] = _format_number(numbers[0], 1)
            payload["pulmonary_backward_volume"] = _format_number(numbers[1], 1)
        elif re.match(r"^regurgitant fraction\b", lowered):
            payload["ar_rf"] = _format_number(numbers[0], 1)
            payload["pr_rf"] = _format_number(numbers[1], 1)


def _cmr_valve_severity(text: str, valve_terms: list[str], lesion_terms: list[str]) -> str | None:
    """Extract concise valve lesion severity from CMR narrative text."""
    valve = r"|".join(re.escape(term) for term in valve_terms)
    lesion = r"|".join(re.escape(term) for term in lesion_terms)
    severity = r"(trivial|trace|mild(?:\s+to\s+moderate)?|moderate(?:\s+to\s+severe)?|severe|significant)"
    if re.search(rf"\bno\s+significant\b[^\n.]*\b(?:{valve})\b[^\n.]*\b(?:{lesion})\b", text, flags=re.IGNORECASE):
        return "None"
    if re.search(rf"\bno\s+(?:{valve})\s+(?:{lesion})\b", text, flags=re.IGNORECASE):
        return "None"
    for pattern in [
        rf"\b(?P<sev>{severity})\s+(?:\w+\s+){{0,3}}(?:{valve})\s+(?:{lesion})\b",
        rf"\b(?:{valve})\s+(?:{lesion})\s+(?:is\s+)?(?P<sev>{severity})\b",
    ]:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            raw = _normalize_whitespace(match.group("sev"))
            if raw.lower() == "significant":
                return "Significant"
            return raw[:1].upper() + raw[1:].lower()
    return None


def _backfill_cmr_narrative_fields(payload: dict[str, Any], text: str) -> None:
    """Backfill coded CMR fields from explicit narrative statements."""
    lowered = text.lower()

    for key, patterns in {
        "lv_size": [
            (r"\bleft ventricle\b.{0,80}\b(?:normal in size|not[- ]dilated|non[- ]dilated)\b|\blv\b.{0,80}\b(?:normal in size|not[- ]dilated|non[- ]dilated)\b", "Normal"),
            (r"\bleft ventricle\b.{0,80}\bmild(?:ly)? dilat(?:ed|ation)\b", "Mildly dilated"),
            (r"\bleft ventricle\b.{0,80}\bmoderate(?:ly)? dilat(?:ed|ation)\b", "Moderately dilated"),
            (r"\bleft ventricle\b.{0,80}\bsevere(?:ly)? dilat(?:ed|ation)\b", "Severely dilated"),
        ],
        "rv_size": [
            (r"\bright ventricle\b.{0,80}\b(?:normal in size|not[- ]dilated|non[- ]dilated)\b|\brv\b.{0,80}\b(?:normal in size|not[- ]dilated|non[- ]dilated)\b|\bnormal rv size\b", "Normal"),
            (r"\bright ventricle\b.{0,80}\bmild(?:ly)? dilat(?:ed|ation)\b", "Mildly dilated"),
            (r"\bright ventricle\b.{0,80}\bmoderate(?:ly)? dilat(?:ed|ation)\b", "Moderately dilated"),
            (r"\bright ventricle\b.{0,80}\bsevere(?:ly)? dilat(?:ed|ation)\b", "Severely dilated"),
        ],
        "lv_function": [
            (r"\b(?:normal|preserved)\s+lv systolic function\b|\blv\b.{0,80}\bsystolic function\b.{0,40}\b(?:normal|preserved)\b|\bleft ventricular systolic function\b.{0,40}\b(?:normal|preserved)\b", "Normal"),
            (r"\blv\b.{0,80}\bmild(?:ly)? impaired\b", "Mildly impaired"),
            (r"\blv\b.{0,80}\bmoderate(?:ly)? impaired\b", "Moderately impaired"),
            (r"\blv\b.{0,80}\bsevere(?:ly)? impaired\b", "Severely impaired"),
        ],
        "rv_function": [
            (r"\b(?:normal|preserved)\s+rv systolic function\b|\brv\b.{0,80}\bsystolic function\b.{0,40}\b(?:normal|preserved)\b|\bright ventricular systolic function\b.{0,40}\b(?:normal|preserved)\b", "Normal"),
            (r"\brv\b.{0,100}\bborderline reduced global systolic function\b|\bborderline reduced\s+(?:global\s+)?rv systolic function\b", "Borderline Reduced"),
            (r"\brv\b.{0,100}\bmarkedly reduced longitudinal function\b|\bmarkedly reduced longitudinal rv function\b", "Markedly reduced longitudinal function"),
            (r"\brv\b.{0,80}\blongitudinal systolic function\b.{0,30}\breduced\b", "Reduced longitudinal function"),
            (r"\brv\b.{0,80}\bmild(?:ly)? impaired\b", "Mildly impaired"),
            (r"\brv\b.{0,80}\bmoderate(?:ly)? impaired\b", "Moderately impaired"),
            (r"\brv\b.{0,80}\bsevere(?:ly)? impaired\b", "Severely impaired"),
        ],
        "la_size": [
            (r"\bleft (?:atrium|atria)\b.{0,60}\bnormal in size\b", "Normal"),
            (r"\bleft (?:atrium|atria)\b.{0,60}\bdilat(?:ed|ation)\b", "Dilated"),
        ],
        "ra_size": [
            (r"\bright (?:atrium|atria)\b.{0,60}\bnormal in size\b", "Normal"),
            (r"\bright (?:atrium|atria)\b.{0,60}\bdilat(?:ed|ation)\b", "Dilated"),
        ],
    }.items():
        _set_if_missing(payload, key, _find_echo_descriptor(text, patterns))

    if "left and right atria are normal in size" in lowered:
        _set_if_missing(payload, "la_size", "Normal")
        _set_if_missing(payload, "ra_size", "Normal")

    if not _has_meaningful_value(payload.get("lvh")):
        match = re.search(r"\b(mild|moderate|severe)?\s*(eccentric|concentric)?\s*hypertrophy\b", text, flags=re.IGNORECASE)
        if match:
            payload["lvh"] = _normalize_whitespace(match.group(0)).capitalize()

    if not _has_meaningful_value(payload.get("rwma")):
        if re.search(r"\bno regional wall motion abnormality\b", text, flags=re.IGNORECASE):
            payload["rwma"] = "No"
        elif re.search(r"\bregional wall motion abnormality\b", text, flags=re.IGNORECASE):
            payload["rwma"] = "Yes"

    if not _has_meaningful_value(payload.get("d_shaped_lv")):
        if re.search(r"\bno interventricular septal flattening\b|\bseptal motion is normal\b", text, flags=re.IGNORECASE):
            payload["d_shaped_lv"] = "No"
        elif re.search(r"\bd[- ]shaped\b|\bseptal flattening\b", text, flags=re.IGNORECASE):
            payload["d_shaped_lv"] = "Yes"
    if not _has_meaningful_value(payload.get("septal_flattening")):
        if re.search(r"\bno interventricular septal flattening\b|\bseptal motion is normal\b", text, flags=re.IGNORECASE):
            payload["septal_flattening"] = "Absent"
        elif re.search(r"\bseptal flattening\b|\bflattened septum\b", text, flags=re.IGNORECASE):
            payload["septal_flattening"] = "Present"
    if not _has_meaningful_value(payload.get("flattening_phase")) and re.search(
        r"\bsystolic septal flattening\b|\bseptal flattening\b.{0,30}\bsystol",
        text,
        flags=re.IGNORECASE,
    ):
        payload["flattening_phase"] = "Systolic"

    if not _has_meaningful_value(payload.get("ias_bowing")) and re.search(r"interatrial septum bows", text, flags=re.IGNORECASE):
        payload["ias_bowing"] = "Present"
    if not _has_meaningful_value(payload.get("ias_direction")):
        if re.search(r"interatrial septum bows towards the left atrium", text, flags=re.IGNORECASE):
            payload["ias_direction"] = "Towards left atrium"
        elif re.search(r"interatrial septum bows towards the right atrium", text, flags=re.IGNORECASE):
            payload["ias_direction"] = "Towards right atrium"
    if not _has_meaningful_value(payload.get("rap")) and re.search(r"elevated right atrial pressure", text, flags=re.IGNORECASE):
        payload["rap"] = "Elevated"

    vortex_match = re.search(
        r"vortex formation is present(?: in the main pulmonary artery)?(?:\s*\(([^)]*)\))?",
        text,
        flags=re.IGNORECASE,
    )
    if vortex_match:
        suffix = f" ({_normalize_whitespace(vortex_match.group(1))})" if vortex_match.group(1) else ""
        _set_if_missing(payload, "mpa_vortex", f"Present{suffix}")

    if not _has_meaningful_value(payload.get("lge")):
        if re.search(r"\bno (?:late gadolinium enhancement|lge|scar|fibrosis)\b", text, flags=re.IGNORECASE):
            payload["lge"] = "Absent"
        elif re.search(r"\blate gadolinium enhancement\b|\blge\b|\bsubendocardial\b|\btransmural\b|\bmid[- ]wall\b", text, flags=re.IGNORECASE):
            payload["lge"] = "Present"
    if not _has_meaningful_value(payload.get("fibrosis")) and re.search(r"\bfibrosis\b", text, flags=re.IGNORECASE):
        payload["fibrosis"] = "Present"
    if re.search(r"\bsubendocardial\b", text, flags=re.IGNORECASE):
        _set_if_missing(payload, "lge_pattern", "Subendocardial")
    if re.search(r"\bmid[- ]wall\b", text, flags=re.IGNORECASE):
        _set_if_missing(payload, "lge_pattern", "Mid-wall")
    if re.search(r"\btransmural\b", text, flags=re.IGNORECASE):
        _set_if_missing(payload, "lge_transmurality", "Transmural")
    lge_location_match = re.search(
        r"(?:enhancement|lge).{0,60}\bin the ([^.:\n;]+?)(?:\.|;|\n|$)",
        text,
        flags=re.IGNORECASE,
    )
    if lge_location_match:
        _set_if_missing(payload, "lge_location", _normalize_whitespace(lge_location_match.group(1)))
    if re.search(r"\b(rv insertion[- ]point fibrosis|right ventricular insertion[- ]point fibrosis)\b", text, flags=re.IGNORECASE):
        _set_if_missing(payload, "rv_insertion_point_lge", "Present")
    if not _has_meaningful_value(payload.get("mpa_flow")) and re.search(
        r"\b(?:disorgani[sz]ed|helical|vertical)\b.{0,80}\bflow\b|\bflow\b.{0,80}\b(?:disorgani[sz]ed|helical|vertical)\b",
        text,
        flags=re.IGNORECASE,
    ):
        payload["mpa_flow"] = "Disorganised vertical/helical flow"

    for key, valve_terms, lesion_terms in [
        ("mr_severity", ["mitral", "MR"], ["regurgitation", "MR"]),
        ("tr_severity", ["tricuspid", "TR"], ["regurgitation", "TR"]),
        ("pr_severity", ["pulmonary", "pulmonic", "PR"], ["regurgitation", "PR"]),
        ("ar_severity", ["aortic", "AR"], ["regurgitation", "AR"]),
        ("as_severity", ["aortic", "AS"], ["stenosis", "AS"]),
    ]:
        _set_if_missing(payload, key, _cmr_valve_severity(text, valve_terms, lesion_terms))
    if re.search(r"\bno significant aortic or tricuspid regurgitation\b", text, flags=re.IGNORECASE):
        _set_if_missing(payload, "ar_severity", "None")
        _set_if_missing(payload, "tr_severity", "None")

    if not _has_meaningful_value(payload.get("pericardial_effusion")):
        if re.search(r"\bno pericardial effusion\b", text, flags=re.IGNORECASE):
            payload["pericardial_effusion"] = "Absent"
        elif re.search(r"\bpericardial effusion\b", text, flags=re.IGNORECASE):
            payload["pericardial_effusion"] = "Present"


def _backfill_cmr_from_text(
    payload: dict[str, Any],
    source_text: str | None,
) -> None:
    """Apply deterministic CMR backfills from source text and derived measurements."""
    text = source_text or ""

    normalized_date = _normalize_date_string(payload.get("date_cmr"))
    if normalized_date:
        payload["date_cmr"] = normalized_date

    if not text:
        return

    if not _has_meaningful_value(payload.get("date_cmr")):
        _set_if_missing(payload, "date_cmr", _extract_cmr_date_from_text(text))
    _set_if_missing(payload, "indication", _extract_indication_from_text(text))
    _set_if_missing(payload, "conclusions", _extract_cmr_conclusions(text))

    # Normalize heart-rate ranges to their arithmetic mean.
    heart_rate_match = re.search(
        r"heart rate(?:\s+was)?(?:\s+between)?\s+(\d+(?:\.\d+)?)\s*(?:-|to|–|—)\s*(\d+(?:\.\d+)?)\s*bpm",
        text,
        flags=re.IGNORECASE,
    )
    if heart_rate_match:
        mean_hr = (float(heart_rate_match.group(1)) + float(heart_rate_match.group(2))) / 2
        _set_if_missing(payload, "heart_rate", _format_number(mean_hr, 1), overwrite=True)
    else:
        single_hr_match = re.search(
            r"heart rate(?:\s+was)?(?:\s+of)?\s+(\d+(?:\.\d+)?)\s*bpm",
            text,
            flags=re.IGNORECASE,
        )
        if single_hr_match:
            _set_if_missing(payload, "heart_rate", _format_number(float(single_hr_match.group(1)), 1))

    robust_hr_range = re.search(
        r"heart rate(?:\s+was)?(?:\s+between)?[^\d]{0,20}(\d+(?:\.\d+)?)\s*(?:-|to|\u2013|\u2014)\s*(\d+(?:\.\d+)?)\s*(?:bpm)?",
        text,
        flags=re.IGNORECASE,
    )
    if robust_hr_range:
        mean_hr = (float(robust_hr_range.group(1)) + float(robust_hr_range.group(2))) / 2
        _set_if_missing(payload, "heart_rate", _format_number(mean_hr, 1), overwrite=True)

    average_hr_match = re.search(
        r"heart rate[^\n.]{0,60}\b(?:averaged|average(?:d)?(?:\s+of|\s+was)?)\s+(\d+(?:\.\d+)?)\s*bpm",
        text,
        flags=re.IGNORECASE,
    )
    if average_hr_match:
        _set_if_missing(
            payload,
            "heart_rate",
            _format_number(float(average_hr_match.group(1)), 1),
            overwrite=True,
        )

    cmr_numeric_specs: dict[str, tuple[list[str], int]] = {
        "height": (["height"], 1),
        "weight": (["weight"], 1),
        "lvedv": (["LV EDV (mL)", "LVEDV"], 1),
        "lvesv": (["LV ESV (mL)", "LVESV"], 1),
        "lvsv": (["LV SV (mL)", "LVSV"], 1),
        "lvedvi": (["LV EDV (i)", "LVEDVi", "LV EDV index"], 1),
        "lvesvi": (["LV ESV (i)", "LVESVi", "LV ESV index"], 1),
        "lvsvi": (["LV SV (i)", "LVSVi", "LV SV index"], 1),
        "lv_mass": (["LV mass (g)", "LV mass"], 1),
        "lvmi": (["LV mass (i)", "LVMi", "LV mass index"], 1),
        "max_lv_wall": (["LV peak wall thickness", "max LV wall", "maximum LV wall"], 1),
        "lvef": (["LV EF", "LVEF"], 0),
        "rvedv": (["RV EDV (mL)", "RVEDV"], 1),
        "rvesv": (["RV ESV (mL)", "RVESV"], 1),
        "rvsv": (["RV SV (mL)", "RVSV"], 1),
        "rvedvi": (["RV EDV (i)", "RVEDVi", "RV EDV index"], 1),
        "rvesvi": (["RV ESV (i)", "RVESVi", "RV ESV index"], 1),
        "rvsvi": (["RV SV (i)", "RVSVi", "RV SV index"], 1),
        "rvef": (["RV EF", "RVEF"], 0),
        "pcwp": (["Estimated PCWP", "PCWP"], 1),
        "rap": (["Estimated RAP", "RAP"], 1),
        "ao_forward_volume": (["Aortic forward flow", "Aortic forward volume"], 1),
        "ao_backward_volume": (["Aortic backward flow", "Aortic backward volume"], 1),
        "pulmonary_forward_volume": (["Pulmonary forward flow", "Pulmonary forward volume"], 1),
        "pulmonary_backward_volume": (["Pulmonary backward flow", "Pulmonary backward volume"], 1),
        "mpa_size": (["MPA systolic diameter", "MPA diameter", "main pulmonary artery diameter"], 1),
        "native_t1": (["Native myocardial T1", "Native T1"], 0),
        "t2": (["Native myocardial T2", "Native T2"], 0),
        "t2_star": (["T2 star", "T2*"], 1),
        "ecv": (["ECV"], 0),
        "tapse": (["TAPSE"], 1),
        "mapse": (["MAPSE"], 1),
    }
    source_numeric_keys: set[str] = set()
    for key, (labels, decimals) in cmr_numeric_specs.items():
        match = _first_number_after_labels(text, labels)
        if not match:
            continue
        number, _fragment = match
        source_numeric_keys.add(key)
        _set_if_missing(payload, key, _format_number(number, decimals), overwrite=True)

    if re.search(r"\b4d\s*[- ]?flow\b", text, flags=re.IGNORECASE):
        _set_if_missing(payload, "flow", "4D-flow", overwrite=True)
    elif re.search(r"\bflow\s*\(\s*2d\s*-?\s*pc\s*\)", text, flags=re.IGNORECASE):
        _set_if_missing(payload, "flow", "2D-flow", overwrite=True)

    if re.search(r"late gadolinium enhancement|\bgadolinium\b|\blge\b", text, flags=re.IGNORECASE):
        _set_if_missing(payload, "contrast", "Gadolinium")

    tapse_match = re.search(r"\bTAPSE\s*(\d+(?:\.\d+)?)\s*mm\b", text, flags=re.IGNORECASE)
    if tapse_match:
        _set_if_missing(payload, "tapse", _format_number(float(tapse_match.group(1)), 1))

    if re.search(r"left and right atria are normal in size", text, flags=re.IGNORECASE):
        _set_if_missing(payload, "la_size", "Normal")
        _set_if_missing(payload, "ra_size", "Normal")
    else:
        if re.search(r"left atri(?:um|a)\s+(?:is|are)\s+normal in size", text, flags=re.IGNORECASE):
            _set_if_missing(payload, "la_size", "Normal")
        if re.search(r"right atri(?:um|a)\s+(?:is|are)\s+normal in size", text, flags=re.IGNORECASE):
            _set_if_missing(payload, "ra_size", "Normal")

    if re.search(
        r"low probability of significant pulmonary hypertension on cmr|no cmr features to support ph physiology more than mild",
        text,
        flags=re.IGNORECASE,
    ):
        _set_if_missing(payload, "ph", "Low", overwrite=True)
    elif re.search(
        r"(?:cmr (?:features|findings)|features).{0,80}support(?:s)?\s+(?:pulmonary hypertension|ph)\s+physiology|support ph physiology",
        text,
        flags=re.IGNORECASE,
    ):
        _set_if_missing(payload, "ph", "High", overwrite=True)

    _backfill_cmr_narrative_fields(payload, text)
    _backfill_cmr_2d_flow_table(payload, text)

    bsa = _compute_bsa_from_payload(payload)
    if bsa:
        for abs_key, idx_key, decimals in [
            ("lvedv", "lvedvi", 1),
            ("lvesv", "lvesvi", 1),
            ("lvsv", "lvsvi", 1),
            ("lv_mass", "lvmi", 1),
            ("rvedv", "rvedvi", 1),
            ("rvesv", "rvesvi", 1),
            ("rvsv", "rvsvi", 1),
        ]:
            abs_value = _to_float(payload.get(abs_key))
            idx_value = _to_float(payload.get(idx_key))
            has_source_abs = abs_key in source_numeric_keys
            has_source_idx = idx_key in source_numeric_keys
            if idx_value is not None and not has_source_abs:
                _set_if_missing(
                    payload,
                    abs_key,
                    _format_number(idx_value * bsa, decimals),
                    overwrite=True,
                )
            elif abs_value is None and idx_value is not None:
                _set_if_missing(payload, abs_key, _format_number(idx_value * bsa, decimals))
            if abs_value is not None and not has_source_idx:
                _set_if_missing(
                    payload,
                    idx_key,
                    _format_number(abs_value / bsa, decimals),
                    overwrite=has_source_abs,
                )

        la_index_match = re.search(
            r"\bLA(?:\s+max)?\s+volume\s*(?:\(i\))?\s*\(mL/m2\)\s+(\d+(?:\.\d+)?)\b",
            text,
            flags=re.IGNORECASE,
        )
        if la_index_match:
            la_index = float(la_index_match.group(1))
            _set_if_missing(payload, "la_volume", _format_number(la_index * bsa, 1), overwrite=True)
            if la_index <= 41:
                _set_if_missing(payload, "la_size", "Normal")

        ra_index_match = re.search(
            r"\bRA(?:\s+max)?\s+volume\s*(?:\(i\))?\s*\(mL/m2\)\s+(\d+(?:\.\d+)?)\b",
            text,
            flags=re.IGNORECASE,
        )
        if ra_index_match and float(ra_index_match.group(1)) <= 48:
            _set_if_missing(payload, "ra_size", "Normal")

    rv_index = _to_float(payload.get("rvedvi"))
    lv_index = _to_float(payload.get("lvedvi"))
    has_source_ratio_inputs = "rvedvi" in source_numeric_keys and "lvedvi" in source_numeric_keys
    if rv_index is not None and lv_index not in (None, 0):
        if has_source_ratio_inputs or not _has_meaningful_value(payload.get("rv_lv_ratio")):
            payload["rv_lv_ratio"] = _format_number(rv_index / lv_index, 2)
    elif not _has_meaningful_value(payload.get("rv_lv_ratio")):
        rv_abs = _to_float(payload.get("rvedv"))
        lv_abs = _to_float(payload.get("lvedv"))
        if rv_abs is not None and lv_abs not in (None, 0):
            payload["rv_lv_ratio"] = _format_number(rv_abs / lv_abs, 2)


def _normalize_extracted_payload(
    modality: str,
    extracted: dict[str, Any],
    source_text: str | None = None,
) -> dict[str, Any]:
    """Map common alias keys onto the record schema expected by the UI and DB."""
    normalized = dict(extracted)

    alias_map: dict[str, str] = {}
    if modality == "rhc":
        alias_map = RHC_ALIAS_MAP
    elif modality == "echo":
        alias_map = ECHO_ALIAS_MAP
    elif modality == "cmr":
        alias_map = {
            "study_date": "date_cmr",
            "date": "date_cmr",
            "patient_class": "cmr_class",
            "case_type": "cmr_class",
            "primary_diagnosis": "primary_dx",
            "secondary_diagnosis": "secondary_dx",
            "conclusion_text_exact": "conclusions",
            "heart_rate_bpm": "heart_rate",
            "lv_function_description": "lv_function",
            "rv_function_description": "rv_function",
            "pulmonary_hypertension_probability": "ph",
        }

    _apply_aliases(normalized, alias_map)

    for date_key in (
        "date_rhc",
        "study_date",
        "report_date",
        "date_cmr",
        "date_cpex",
        "date_of_birth",
    ):
        normalized_date = _normalize_date_string(normalized.get(date_key))
        if normalized_date:
            normalized[date_key] = normalized_date

    if modality == "rhc":
        _normalize_rhc_payload(normalized)
    elif modality == "echo":
        _normalize_echo_payload(normalized)
        _backfill_echo_from_text(normalized, source_text)

    if modality == "cmr" and not _has_meaningful_value(normalized.get("conclusions")):
        conclusion_items = normalized.get("conclusion_items")
        if isinstance(conclusion_items, list):
            conclusion_text = "\n".join(
                str(item).strip() for item in conclusion_items if str(item).strip()
            )
            if conclusion_text:
                normalized["conclusions"] = conclusion_text
        elif isinstance(conclusion_items, str) and conclusion_items.strip():
            normalized["conclusions"] = conclusion_items.strip()

    if modality == "cmr" and not _has_meaningful_value(normalized.get("qc_notes")):
        qc_parts: list[str] = []
        warnings = normalized.get("extraction_warnings")
        uncertain = normalized.get("uncertain_fields")
        if isinstance(warnings, list):
            qc_parts.extend(str(item).strip() for item in warnings if str(item).strip())
        elif isinstance(warnings, str) and warnings.strip():
            qc_parts.append(warnings.strip())
        if isinstance(uncertain, list):
            uncertain_text = ", ".join(
                str(item).strip() for item in uncertain if str(item).strip()
            )
            if uncertain_text:
                qc_parts.append(f"Uncertain fields: {uncertain_text}")
        elif isinstance(uncertain, str) and uncertain.strip():
            qc_parts.append(f"Uncertain fields: {uncertain.strip()}")
        if qc_parts:
            normalized["qc_notes"] = "\n".join(qc_parts)

    if modality == "cmr":
        _backfill_cmr_from_text(normalized, source_text)
        _normalize_cmr_coded_values(normalized, source_text)
        derived_class = _derive_cmr_class(
            str(normalized.get("primary_dx") or ""),
            str(normalized.get("conclusions") or ""),
        )
        _set_if_missing(normalized, "cmr_class", derived_class)
        if not _has_meaningful_value(normalized.get("qc_notes")):
            fallback_qc_parts: list[str] = []
            if not _has_meaningful_value(normalized.get("indication")):
                fallback_qc_parts.append(
                    "Review extraction: indication was not clearly captured."
                )
            primary_dx = str(normalized.get("primary_dx") or "").strip().lower()
            cmr_class = str(normalized.get("cmr_class") or "").strip().lower()
            conclusions = str(normalized.get("conclusions") or "").strip().lower()
            has_positive_ischaemic_primary = any(
                term in primary_dx
                for term in [
                    "prior myocardial infarction",
                    "ischaemic scar",
                    "ischemic scar",
                    "ischaemic cardiomyopathy",
                    "ischemic cardiomyopathy",
                ]
            ) and "non-ischaemic" not in primary_dx and "non-ischemic" not in primary_dx
            if (
                has_positive_ischaemic_primary
                and (
                    "no evidence of any scar" in conclusions
                    or "no scar" in conclusions
                    or "no fibrosis" in conclusions
                    or "no lge" in conclusions
                )
            ):
                fallback_qc_parts.append(
                    "Review classification: conclusions describe no scar/fibrosis despite an ischaemic primary diagnosis."
                )
            if (
                has_positive_ischaemic_primary
                and cmr_class
                and any(
                    token in cmr_class
                    for token in ["non-ischaemic", "myocarditis", "inflammatory"]
                )
            ):
                fallback_qc_parts.append(
                    "Review classification: primary diagnosis and CMR class may be inconsistent."
                )
            if fallback_qc_parts:
                normalized["qc_notes"] = "\n".join(fallback_qc_parts)

    return normalized


def _coerce_record_value_for_storage(value: Any) -> Any:
    """Coerce extracted complex values into scalar DB-safe representations."""
    if value is None:
        return None
    if isinstance(value, (list, tuple, set)):
        scalar_items = [
            str(item).strip()
            for item in value
            if not isinstance(item, (list, tuple, set, dict)) and _has_meaningful_value(item)
        ]
        if scalar_items and len(scalar_items) == len(value):
            return "\n".join(scalar_items)
        return json.dumps(value, ensure_ascii=False)
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    return value


def _coerce_numeric_record_value(value: Any, *, integer: bool = False) -> float | int | None:
    """Coerce extracted numeric column values without letting labels break DB inserts."""
    if value is None:
        return None
    if isinstance(value, bool):
        return int(value) if integer else None
    if isinstance(value, (int, float)):
        if not math.isfinite(float(value)):
            return None
        return int(round(float(value))) if integer else float(value)
    if isinstance(value, str):
        text = value.strip()
        if not text or text in {"-", "\u2014", "\u2013"}:
            return None
        lowered = text.lower()
        if integer:
            if lowered in {"yes", "true", "present", "positive"}:
                return 1
            if lowered in {"no", "false", "absent", "negative", "none", "nil"}:
                return 0
        number = _to_float(text)
        if number is None:
            return None
        return int(round(number)) if integer else number
    return None


def _coerce_record_data_for_storage(
    modality: str,
    record_data: dict[str, Any],
    records_svc: Any,
) -> dict[str, Any]:
    """Convert extracted payload values into DB-safe scalars using model column types."""
    model_cls = records_svc.MODALITY_MAP[modality][0]
    columns = model_cls.__table__.columns
    coerced: dict[str, Any] = {}
    for key, value in record_data.items():
        column = columns.get(key)
        if column is not None and isinstance(column.type, Integer):
            coerced[key] = _coerce_numeric_record_value(value, integer=True)
            continue
        if column is not None and isinstance(column.type, Float):
            coerced[key] = _coerce_numeric_record_value(value)
            continue
        coerced[key] = _coerce_record_value_for_storage(value)
    return coerced


def _build_text_messages_for_modality(
    prompt: str,
    text: str,
    modality: str,
) -> list[dict[str, str]]:
    """Build chat messages for text extraction with an explicit JSON contract."""
    schema_instruction = _schema_instruction(modality)
    system_content = prompt.rstrip()
    if schema_instruction:
        system_content = f"{system_content}\n\n{schema_instruction}"
    system_content = (
        f"{system_content}\n\n"
                "Return a valid JSON object only. Do not wrap the JSON in markdown fences."
    )
    return [
        {
            "role": "system",
            "content": system_content,
        },
        {
            "role": "user",
            "content": (
                "Extract the structured data from the following clinical text and "
                "return a JSON object only.\n\n"
                f"{text}"
            ),
        },
    ]


# ---------------------------------------------------------------------------
# Core extraction functions
# ---------------------------------------------------------------------------


def extract_from_text(
    text: str,
    modality: str,
    source_type: str | None = None,
) -> dict[str, Any]:
    """Send text to GPT-4o with the appropriate prompt. Return extracted data."""
    prompt = _get_prompt(modality, source_type)
    client = _get_client()
    model = _get_model()
    messages = _build_text_messages_for_modality(prompt, text, modality)

    try:
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            response_format={"type": "json_object"},
            temperature=0.0,
        )
    except openai.APITimeoutError:
        raise ExtractionTimeoutError("OpenAI API request timed out after 120 seconds")
    except openai.RateLimitError:
        raise ExtractionRateLimitError("OpenAI API rate limit exceeded. Please retry later.")
    except openai.BadRequestError as exc:
        if "maximum context length" in str(exc).lower():
            raise ExtractionTokenLimitError(
                "Document exceeds the model's maximum token limit. Try a shorter document."
            )
        raise

    raw_content = response.choices[0].message.content or ""
    logger.info(
        "Text extraction response length=%d, finish_reason=%s",
        len(raw_content),
        response.choices[0].finish_reason,
    )

    try:
        extracted = _normalize_extracted_payload(
            modality, _parse_json_response(raw_content), text
        )
    except (json.JSONDecodeError, ValueError):
        logger.warning("First extraction returned malformed JSON (len=%d), retrying once. Raw: %s", len(raw_content), raw_content[:200])
        try:
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                response_format={"type": "json_object"},
                temperature=0.0,
            )
            raw_content = response.choices[0].message.content or ""
            extracted = _normalize_extracted_payload(
                modality, _parse_json_response(raw_content), text
            )
        except (json.JSONDecodeError, ValueError, openai.APITimeoutError, openai.RateLimitError):
            raise ExtractionParseError(
                f"GPT-4o returned malformed JSON after retry. Raw output: {raw_content[:500]}"
            )

    nonempty_keys = [key for key, value in extracted.items() if _has_meaningful_value(value)]
    logger.info(
        "Parsed extraction modality=%s keys=%d nonempty=%d sample_keys=%s",
        modality,
        len(extracted),
        len(nonempty_keys),
        nonempty_keys[:12],
    )

    return {
        "modality": modality,
        "source_type": source_type,
        "extracted_data": extracted,
    }


def extract_from_image(
    image_base64: str,
    modality: str,
    source_type: str | None = None,
) -> dict[str, Any]:
    """Send an image to GPT-4o via vision API and extract data."""
    prompt = _get_prompt(modality, source_type)
    client = _get_client()
    model = _get_model()

    # Determine mime type from base64 header or default to png
    if image_base64.startswith("/9j/"):
        mime = "image/jpeg"
    else:
        mime = "image/png"

    data_uri = f"data:{mime};base64,{image_base64}"
    schema_instruction = _schema_instruction(modality)
    image_prompt = prompt
    if schema_instruction:
        image_prompt = f"{image_prompt}\n\n{schema_instruction}"
    image_prompt = (
        image_prompt
        + "\n\nExtract the structured data from the attached clinical document image. "
        + "You MUST return a JSON object with all extracted fields."
    )

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": image_prompt,
                        },
                        {
                            "type": "image_url",
                            "image_url": {"url": data_uri, "detail": "high"},
                        },
                    ],
                },
            ],
            response_format={"type": "json_object"},
            temperature=0.0,
            max_tokens=4096,
        )
    except openai.APITimeoutError:
        raise ExtractionTimeoutError("OpenAI API request timed out after 120 seconds")
    except openai.RateLimitError:
        raise ExtractionRateLimitError("OpenAI API rate limit exceeded. Please retry later.")
    except openai.BadRequestError as exc:
        if "maximum context length" in str(exc).lower():
            raise ExtractionTokenLimitError(
                "Image exceeds the model's maximum token limit."
            )
        raise

    raw_content = response.choices[0].message.content or ""

    try:
        extracted = _normalize_extracted_payload(
            modality, _parse_json_response(raw_content)
        )
    except (json.JSONDecodeError, ValueError):
        logger.warning("First image extraction returned malformed JSON, retrying once")
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": image_prompt,
                            },
                            {
                                "type": "image_url",
                                "image_url": {"url": data_uri, "detail": "high"},
                            },
                        ],
                    },
                ],
                response_format={"type": "json_object"},
                temperature=0.0,
                max_tokens=4096,
            )
            raw_content = response.choices[0].message.content or ""
            extracted = _normalize_extracted_payload(
                modality, _parse_json_response(raw_content)
            )
        except (json.JSONDecodeError, ValueError, openai.APITimeoutError, openai.RateLimitError):
            raise ExtractionParseError(
                f"GPT-4o returned malformed JSON after retry. Raw output: {raw_content[:500]}"
            )

    return {
        "modality": modality,
        "source_type": source_type,
        "extracted_data": extracted,
    }


def extract_from_file(
    file_bytes: bytes,
    filename: str,
    modality: str,
    source_type: str | None = None,
) -> dict[str, Any]:
    """Extract text from a file, then call the appropriate extraction function."""
    if len(file_bytes) > MAX_FILE_SIZE:
        raise ExtractionFileTooLargeError(
            f"File exceeds the 20 MB limit ({len(file_bytes)} bytes)"
        )

    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        return extract_from_text(text, modality, source_type)

    elif ext == ".docx":
        text = _extract_docx_text(file_bytes)
        if not text.strip():
            raise ExtractionParseError(
                "No readable text was found in the DOCX file. Try saving it as PDF or uploading a screenshot."
            )
        return extract_from_text(text, modality, source_type)

    elif ext == ".doc":
        text = _extract_doc_text(file_bytes)
        return extract_from_text(text, modality, source_type)

    elif ext in (".png", ".jpg", ".jpeg"):
        b64 = base64.b64encode(file_bytes).decode()
        return extract_from_image(b64, modality, source_type)

    else:
        raise ExtractionUnsupportedFileError(f"Unsupported file type: {ext}")


# ---------------------------------------------------------------------------
# Save extraction results
# ---------------------------------------------------------------------------


def save_extraction(
    modality: str,
    hospital_number: str,
    create_patient_if_missing: bool,
    patient_data: dict[str, Any],
    record_data: dict[str, Any],
    source_file_upload_id: str | None = None,
) -> dict[str, Any]:
    """Save extraction result to database via patients and records services."""
    from research_os.extract_patients import service as patients_svc
    from research_os.extract_records import service as records_svc
    from research_os.extract_source_files import service as source_files_svc

    if create_patient_if_missing:
        patient = patients_svc.find_or_create_patient(hospital_number, **patient_data)
    else:
        patient = patients_svc.get_patient(hospital_number)

    if modality in {"rhc", "echo"}:
        record_data = _normalize_extracted_payload(modality, record_data)

    # Only keep keys that are valid columns on the record model — GPT-4o
    # returns many extra keys (patient-level fields, warnings, etc.)
    valid_columns = records_svc.get_valid_columns(modality)
    record_data = {k: v for k, v in record_data.items() if k in valid_columns}
    source_file_upload: dict[str, Any] | None = None
    clean_source_file_upload_id = str(source_file_upload_id or "").strip()
    if clean_source_file_upload_id:
        source_file_upload = source_files_svc.get_source_file(clean_source_file_upload_id)
        if source_file_upload.get("modality") != modality:
            raise ValueError("Source file modality does not match extraction modality")
        if "source_file" in valid_columns and not record_data.get("source_file"):
            record_data["source_file"] = source_file_upload.get("filename")
    record_data = _coerce_record_data_for_storage(modality, record_data, records_svc)
    if "status" in valid_columns and not record_data.get("status"):
        record_data["status"] = "Completed"
    if (
        record_data.get("status") == "Completed"
        and "status_date" in valid_columns
        and not record_data.get("status_date")
    ):
        record_data["status_date"] = datetime.now(timezone.utc).date().isoformat()
    record_data["hn"] = hospital_number
    record = records_svc.create_record(modality, record_data)
    if clean_source_file_upload_id:
        source_file_upload = source_files_svc.link_source_file(
            file_id=clean_source_file_upload_id,
            modality=modality,
            hn=hospital_number,
            record_id=str(record["id"]),
        )
        record["source_file_upload"] = source_file_upload

    return {"patient": patient, "record": record}


# ---------------------------------------------------------------------------
# Exception classes
# ---------------------------------------------------------------------------


class ExtractionError(RuntimeError):
    """Base class for extraction errors."""
    pass


class ExtractionTimeoutError(ExtractionError):
    """OpenAI API timed out."""
    pass


class ExtractionRateLimitError(ExtractionError):
    """OpenAI rate limit hit."""
    pass


class ExtractionParseError(ExtractionError):
    """GPT-4o returned malformed JSON."""
    pass


class ExtractionTokenLimitError(ExtractionError):
    """Document exceeded model token limit."""
    pass


class ExtractionFileTooLargeError(ExtractionError):
    """Uploaded file exceeds size limit."""
    pass


class ExtractionUnsupportedFileError(ExtractionError):
    """Unsupported file type."""
    pass
